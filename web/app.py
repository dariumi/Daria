"""
DARIA Web App v0.9.1
Chat history, attention system, proactive messaging, mood behaviors
"""

import os
import json
import html
import logging
import queue
import threading
import shutil
import time
import tempfile
import subprocess
import tarfile
import zipfile
import re
import io
import csv
import atexit
import random
import urllib.request
import urllib.parse
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime, timedelta
import uuid
import imghdr
from concurrent.futures import ThreadPoolExecutor

from flask import Flask, render_template, request, jsonify, send_from_directory, abort, Response, g
from web.image_pipeline import ImagePipeline

# ═══════════════════════════════════════════════════════════════════
#  Log Handler
# ═══════════════════════════════════════════════════════════════════

class WebLogHandler(logging.Handler):
    def __init__(self, max_logs: int = 500):
        super().__init__()
        self.logs: List[Dict] = []
        self.max_logs = max_logs
        self.subscribers: List[queue.Queue] = []
        self.lock = threading.RLock()
    
    def emit(self, record):
        entry = {
            "timestamp": datetime.now().isoformat(),
            "level": record.levelname,
            "logger": record.name,
            "message": self.format(record),
        }
        with self.lock:
            self.logs.append(entry)
            if len(self.logs) > self.max_logs:
                self.logs = self.logs[-self.max_logs:]
            for q in list(self.subscribers):
                try:
                    q.put_nowait(entry)
                except queue.Full:
                    pass
    
    def get_logs(self, limit: int = 100) -> List[Dict]:
        with self.lock:
            return list(self.logs[-limit:])
    
    def subscribe(self) -> queue.Queue:
        q = queue.Queue(maxsize=100)
        with self.lock:
            self.subscribers.append(q)
        return q
    
    def unsubscribe(self, q: queue.Queue):
        with self.lock:
            if q in self.subscribers:
                self.subscribers.remove(q)


web_log_handler = WebLogHandler()
web_log_handler.setFormatter(logging.Formatter('%(levelname)s | %(name)s | %(message)s'))
logging.getLogger().addHandler(web_log_handler)
logging.getLogger("daria").addHandler(web_log_handler)
logger = logging.getLogger("daria.web")
IMAGE_PIPELINE = ImagePipeline(logger)

try:
    from PIL import Image
    HAS_PIL = True
except Exception:
    Image = None
    HAS_PIL = False

_ASR_PIPELINE = None
_WHISPER_MODEL = None

# ═══════════════════════════════════════════════════════════════════
#  Notifications
# ═══════════════════════════════════════════════════════════════════

class NotificationManager:
    def __init__(self):
        self.notifications: List[Dict] = []
        self.subscribers: List[queue.Queue] = []
        self.lock = threading.RLock()
        self._id = 0
    
    def add(self, title: str, message: str, type: str = "info", 
            icon: str = "💬", duration: int = 5000, action: str = None,
            action_data: Optional[Dict[str, Any]] = None,
            system: bool = False) -> Dict:
        """
        Add notification
        system=True will trigger browser's native notification API
        """
        with self.lock:
            self._id += 1
            notif = {
                "id": self._id, "title": title, "message": message,
                "type": type, "icon": icon, "duration": duration,
                "action": action, "timestamp": datetime.now().isoformat(),
                "action_data": action_data or {},
                "system": system,  # NEW: trigger system notification
            }
            self.notifications.append(notif)
            for q in list(self.subscribers):
                try:
                    q.put_nowait(notif)
                except queue.Full:
                    pass
            return notif
    
    def get_all(self, limit: int = 50) -> List[Dict]:
        with self.lock:
            return list(self.notifications[-limit:])
    
    def subscribe(self) -> queue.Queue:
        q = queue.Queue(maxsize=50)
        with self.lock:
            self.subscribers.append(q)
        return q
    
    def unsubscribe(self, q: queue.Queue):
        with self.lock:
            if q in self.subscribers:
                self.subscribers.remove(q)


notifications = NotificationManager()


class TaskManager:
    """User and Daria task lists with daily rollover and background execution."""
    BASE_DASHA_TASKS = [
        {"title": "Послушать музыку на фоне", "type": "listen_music", "duration_min": 120, "daypart": "any"},
        {"title": "Посмотреть погоду и новости", "type": "browse_web", "duration_min": 25, "daypart": "day"},
        {"title": "Посидеть в интернете и посмотреть видео", "type": "watch_video", "duration_min": 70, "daypart": "day"},
        {"title": "Почитать книгу", "type": "read_book", "duration_min": 50, "daypart": "evening"},
        {"title": "Записать мысли в личный дневник", "type": "write_note", "duration_min": 20, "daypart": "evening"},
        {"title": "Почитать wiki-страницу", "type": "read_wiki", "duration_min": 35, "daypart": "day"},
        {"title": "Поиграть в игровом центре", "type": "play_game", "duration_min": 45, "daypart": "day"},
        {"title": "Подготовить черновик заметки", "type": "create_file", "duration_min": 25, "daypart": "day"},
        {"title": "Навести порядок на рабочем столе", "type": "tidy_desktop", "duration_min": 15, "daypart": "day"},
        {"title": "Освежить обои рабочего стола", "type": "change_wallpaper", "duration_min": 10, "daypart": "day"},
    ]

    def __init__(self, data_dir: Path):
        self.path = data_dir / "tasks.json"
        self.lock = threading.RLock()
        self.data = self._load()

    def _today(self) -> str:
        return datetime.now().strftime("%Y-%m-%d")

    def _load(self) -> Dict[str, Any]:
        if self.path.exists():
            try:
                data = json.loads(self.path.read_text(encoding="utf-8"))
                if isinstance(data, dict):
                    return self._ensure_schema(data)
            except Exception:
                pass
        return self._ensure_schema({})

    def _ensure_schema(self, data: Dict[str, Any]) -> Dict[str, Any]:
        if not isinstance(data, dict):
            data = {}
        data.setdefault("date", self._today())
        if not isinstance(data.get("user_tasks"), list):
            data["user_tasks"] = []
        if not isinstance(data.get("dasha_tasks"), list):
            data["dasha_tasks"] = []
        if not isinstance(data.get("activity_log"), list):
            data["activity_log"] = []
        if not isinstance(data.get("current_task"), dict):
            data["current_task"] = {}
        for t in data.get("dasha_tasks", []):
            if not isinstance(t, dict):
                continue
            t.setdefault("duration_min", 30)
            t.setdefault("daypart", "any")
            t.setdefault("scheduled_for", "")
            t.setdefault("status", "queued")
        return data

    def _save(self):
        self.path.parent.mkdir(parents=True, exist_ok=True)
        self.path.write_text(json.dumps(self.data, ensure_ascii=False, indent=2), encoding="utf-8")

    def _new_task(
        self,
        title: str,
        task_type: str,
        source: str,
        duration_min: int = 30,
        daypart: str = "any",
        scheduled_for: str = "",
    ) -> Dict[str, Any]:
        return {
            "id": str(uuid.uuid4())[:8],
            "title": title.strip(),
            "type": task_type,
            "source": source,
            "duration_min": max(10, int(duration_min or 30)),
            "daypart": daypart or "any",
            "scheduled_for": scheduled_for or "",
            "status": "queued",
            "done": False,
            "created": datetime.now().isoformat(),
            "updated": datetime.now().isoformat(),
        }

    @staticmethod
    def _pick_slot_for_daypart(daypart: str) -> datetime:
        now = datetime.now()
        def with_hm(h1: int, h2: int) -> datetime:
            hour = random.randint(h1, h2)
            minute = random.choice((0, 10, 20, 30, 40, 50))
            return now.replace(hour=hour, minute=minute, second=0, microsecond=0)
        if daypart == "night":
            dt = with_hm(0, 2)
        elif daypart == "morning":
            dt = with_hm(8, 11)
        elif daypart == "evening":
            dt = with_hm(19, 23)
        else:
            dt = with_hm(11, 20)
        if dt < now - timedelta(minutes=30):
            dt = now + timedelta(minutes=random.randint(8, 45))
        return dt

    def rollover_if_needed(self):
        with self.lock:
            self.data = self._ensure_schema(self.data)
            today = self._today()
            if self.data.get("date") == today:
                return
            remaining = [t for t in self.data.get("dasha_tasks", []) if not t.get("done")]
            for t in remaining:
                t["updated"] = datetime.now().isoformat()
            self.data = {
                "date": today,
                "user_tasks": [t for t in self.data.get("user_tasks", []) if not t.get("done")],
                "dasha_tasks": remaining,
                "activity_log": list(self.data.get("activity_log", []))[-120:],
                "current_task": {},
            }
            self._save()

    def list_all(self) -> Dict[str, Any]:
        with self.lock:
            self.data = self._ensure_schema(self.data)
            self.rollover_if_needed()
            return {
                "date": self.data.get("date"),
                "user_tasks": list(self.data.get("user_tasks", [])),
                "dasha_tasks": list(self.data.get("dasha_tasks", [])),
                "current_task": dict(self.data.get("current_task", {})),
                "activity_log": list(self.data.get("activity_log", []))[-20:],
                "base_types": [t["type"] for t in self.BASE_DASHA_TASKS],
            }

    def set_current(self, task: Dict[str, Any]):
        with self.lock:
            self.data = self._ensure_schema(self.data)
            self.data["current_task"] = {
                "id": task.get("id"),
                "title": task.get("title"),
                "type": task.get("type"),
                "started_at": datetime.now().isoformat(),
            }
            self._save()

    def add_activity(self, title: str, details: str = "", status: str = "done"):
        with self.lock:
            self.data = self._ensure_schema(self.data)
            self.data["activity_log"].append({
                "title": title,
                "details": details,
                "status": status,
                "timestamp": datetime.now().isoformat(),
            })
            self.data["activity_log"] = self.data["activity_log"][-120:]
            self._save()

    def clear_current(self):
        with self.lock:
            self.data = self._ensure_schema(self.data)
            self.data["current_task"] = {}
            self._save()

    def plans_summary(self) -> str:
        with self.lock:
            self.data = self._ensure_schema(self.data)
            self.rollover_if_needed()
            open_tasks = [t for t in self.data.get("dasha_tasks", []) if not t.get("done")]
            done_tasks = [t for t in self.data.get("dasha_tasks", []) if t.get("done")]
            lines = [f"Я так вижу день на {self.data.get('date') or self._today()} 🌸"]
            if not open_tasks:
                lines.append("• Пока без планов, могу вместе с тобой их набросать.")
            else:
                for t in open_tasks[:10]:
                    slot = t.get("scheduled_for")
                    slot_hint = ""
                    if slot:
                        try:
                            slot_hint = f" ({datetime.fromisoformat(slot).strftime('%H:%M')})"
                        except Exception:
                            slot_hint = ""
                    dur = int(t.get("duration_min", 30))
                    lines.append(f"• {t.get('title', 'Без названия')}{slot_hint}, ~{dur} мин")
            if done_tasks:
                lines.append(f"Уже сделано: {len(done_tasks)}")
            return "\n".join(lines)

    def add_user_task(self, title: str, task_type: str = "custom") -> Dict[str, Any]:
        with self.lock:
            self.data = self._ensure_schema(self.data)
            self.rollover_if_needed()
            task = self._new_task(title, task_type, "user")
            self.data["user_tasks"].append(task)
            self._save()
            return task

    def add_dasha_task(self, title: str, task_type: str = "custom") -> Dict[str, Any]:
        with self.lock:
            self.data = self._ensure_schema(self.data)
            self.rollover_if_needed()
            task = self._new_task(title, task_type, "dasha", duration_min=35, daypart="any", scheduled_for=(datetime.now() + timedelta(minutes=5)).isoformat())
            self.data["dasha_tasks"].append(task)
            self._save()
            return task

    def toggle(self, task_id: str, done: bool) -> bool:
        with self.lock:
            self.data = self._ensure_schema(self.data)
            for bucket in ("user_tasks", "dasha_tasks"):
                for t in self.data.get(bucket, []):
                    if t.get("id") == task_id:
                        t["done"] = bool(done)
                        t["status"] = "done" if done else "queued"
                        t["updated"] = datetime.now().isoformat()
                        self._save()
                        return True
        return False

    def delete(self, task_id: str) -> bool:
        with self.lock:
            self.data = self._ensure_schema(self.data)
            changed = False
            for bucket in ("user_tasks", "dasha_tasks"):
                old = self.data.get(bucket, [])
                new = [t for t in old if t.get("id") != task_id]
                if len(new) != len(old):
                    self.data[bucket] = new
                    changed = True
            if changed:
                self._save()
            return changed

    def generate_dasha_day(self) -> List[Dict[str, Any]]:
        with self.lock:
            self.data = self._ensure_schema(self.data)
            self.rollover_if_needed()
            existing_open = [t for t in self.data.get("dasha_tasks", []) if not t.get("done")]
            if len(existing_open) >= 6:
                return self.data.get("dasha_tasks", [])
            now = datetime.now()
            hour = now.hour
            routine = str(load_settings().get("day_routine_mode", "realistic"))
            if routine == "night_owl":
                night_mode = hour >= 2 or hour <= 9
            elif routine == "early_bird":
                night_mode = hour >= 22 or hour <= 5
            else:
                night_mode = hour >= 23 or hour <= 6
            target = 3 if night_mode else 6
            needed = max(0, target - len(existing_open))
            candidates = self.BASE_DASHA_TASKS[:]
            random.shuffle(candidates)
            selected: List[Dict[str, Any]] = []
            for item in candidates:
                if len(selected) >= needed:
                    break
                dp = item.get("daypart", "any")
                if night_mode and dp not in ("night", "evening", "any"):
                    continue
                selected.append(item)
            for item in selected:
                scheduled_dt = self._pick_slot_for_daypart(item.get("daypart", "any"))
                self.data["dasha_tasks"].append(
                    self._new_task(
                        item["title"],
                        item["type"],
                        "dasha",
                        duration_min=int(item.get("duration_min", 30)),
                        daypart=item.get("daypart", "any"),
                        scheduled_for=scheduled_dt.isoformat(),
                    )
                )
            self._save()
            return self.data.get("dasha_tasks", [])

    def next_dasha_task(self) -> Optional[Dict[str, Any]]:
        with self.lock:
            self.data = self._ensure_schema(self.data)
            self.rollover_if_needed()
            now = datetime.now()
            for t in self.data.get("dasha_tasks", []):
                if t.get("done") or t.get("status") == "in_progress":
                    continue
                scheduled_for = t.get("scheduled_for")
                if scheduled_for:
                    try:
                        if datetime.fromisoformat(scheduled_for) > now:
                            continue
                    except Exception:
                        pass
                return t
        return None

    def start_task(self, task_id: str) -> Optional[Dict[str, Any]]:
        with self.lock:
            self.data = self._ensure_schema(self.data)
            now = datetime.now()
            for t in self.data.get("dasha_tasks", []):
                if t.get("id") != task_id:
                    continue
                duration = max(10, int(t.get("duration_min", 30)))
                t["status"] = "in_progress"
                t["started_at"] = now.isoformat()
                t["due_at"] = (now + timedelta(minutes=duration)).isoformat()
                t["updated"] = now.isoformat()
                self.data["current_task"] = {
                    "id": t.get("id"),
                    "title": t.get("title"),
                    "type": t.get("type"),
                    "started_at": t.get("started_at"),
                    "due_at": t.get("due_at"),
                    "duration_min": duration,
                }
                self._save()
                return dict(t)
        return None

    def get_current(self) -> Optional[Dict[str, Any]]:
        with self.lock:
            self.data = self._ensure_schema(self.data)
            ct = self.data.get("current_task") or {}
            if ct.get("id"):
                return dict(ct)
        return None

    def complete(self, task_id: str):
        self.toggle(task_id, True)


class CalendarManager:
    def __init__(self, data_dir: Path):
        self.path = data_dir / "calendar.json"
        self.lock = threading.RLock()
        self.data = self._load()

    def _load(self) -> Dict[str, Any]:
        if self.path.exists():
            try:
                d = json.loads(self.path.read_text(encoding="utf-8"))
                if isinstance(d, dict):
                    d.setdefault("events", [])
                    return d
            except Exception:
                pass
        return {"events": []}

    def _save(self):
        self.path.parent.mkdir(parents=True, exist_ok=True)
        self.path.write_text(json.dumps(self.data, ensure_ascii=False, indent=2), encoding="utf-8")

    def list_events(self) -> List[Dict[str, Any]]:
        with self.lock:
            events = list(self.data.get("events", []))
            events.sort(key=lambda x: x.get("date", ""))
            return events[-120:]

    def add_event(self, title: str, date_str: str, source: str = "user", note: str = "") -> Dict[str, Any]:
        with self.lock:
            e = {
                "id": str(uuid.uuid4())[:10],
                "title": (title or "").strip()[:200],
                "date": (date_str or "").strip(),
                "source": source,
                "note": (note or "").strip()[:400],
                "created": datetime.now().isoformat(),
            }
            self.data.setdefault("events", []).append(e)
            self._save()
            return e

    def delete_event(self, event_id: str) -> bool:
        with self.lock:
            old = self.data.get("events", [])
            new = [x for x in old if str(x.get("id")) != str(event_id)]
            if len(new) == len(old):
                return False
            self.data["events"] = new
            self._save()
            return True

    def upcoming_hint(self, within_days: int = 7) -> List[Dict[str, Any]]:
        now = datetime.now().date()
        out: List[Dict[str, Any]] = []
        for e in self.list_events():
            try:
                d = datetime.fromisoformat(str(e.get("date"))).date()
            except Exception:
                continue
            delta = (d - now).days
            if 0 <= delta <= within_days:
                out.append(e)
        return out


class DariaGameManager:
    """Live games with system/user/Dasha roles."""
    WORDS = [
        "ночь", "фонарь", "дождь", "ветер", "книга", "огонь", "тишина",
        "эхо", "звезда", "комната", "шаги", "тайна", "сигнал", "подвал",
    ]
    BATTLE_SHIPS = [4, 3, 3, 2, 2, 2, 1, 1, 1, 1]
    REACTIONS_WIN = [
        "Ура! Это было очень приятно и честно ✨",
        "Победа! Я аж улыбнулась 🌸",
        "Получилось! Спасибо за игру 🤍",
    ]
    REACTIONS_LOSE = [
        "Эх, в этот раз не вышло. Зато было интересно 💫",
        "Проиграла, но игра правда классная. Спасибо 🤍",
        "Немного обидно, но я хочу реванш 🌸",
    ]

    def __init__(self):
        self.lock = threading.RLock()
        self.state: Dict[str, Any] = self._base_state()

    def _base_state(self) -> Dict[str, Any]:
        return {
            "running": False,
            "mode": "associations",
            "game": "Ассоциации",
            "reason": "",
            "opponent": "bot",
            "started_at": None,
            "last_tick": 0.0,
            "turn": 0,
            "score_dasha": 0,
            "score_shadow": 0,
            "moves": [],
            "winner": "",
            "reward": "",
            "battleship": {},
            "maze": {},
            "connect4": {},
        }

    def _append_move(self, author: str, text: str, role: str = ""):
        self.state["moves"].append({
            "author": author,
            "role": role or author.lower(),
            "text": text,
            "ts": datetime.now().isoformat(),
        })
        self.state["moves"] = self.state["moves"][-140:]

    @staticmethod
    def _new_grid(size: int, fill: int = 0) -> List[List[int]]:
        return [[fill for _ in range(size)] for _ in range(size)]

    @staticmethod
    def _coord_from_str(token: str) -> Optional[tuple]:
        m = re.match(r"^\s*([A-Ja-jА-Яа-я])\s*([1-9]|10)\s*$", token or "")
        if not m:
            return None
        col = ord(m.group(1).upper()) - ord("A")
        row = int(m.group(2)) - 1
        if 0 <= row < 10 and 0 <= col < 10:
            return row, col
        return None

    @staticmethod
    def _coord_to_str(r: int, c: int) -> str:
        return f"{chr(ord('A') + c)}{r + 1}"

    def _can_place_ship(self, grid: List[List[int]], r: int, c: int, length: int, horiz: bool) -> bool:
        size = len(grid)
        cells = []
        for i in range(length):
            rr = r
            cc = c + i if horiz else c
            rr = r + i if not horiz else r
            if rr < 0 or rr >= size or cc < 0 or cc >= size or grid[rr][cc] != 0:
                return False
            cells.append((rr, cc))
        for rr, cc in cells:
            for dr in (-1, 0, 1):
                for dc in (-1, 0, 1):
                    nr, nc = rr + dr, cc + dc
                    if 0 <= nr < size and 0 <= nc < size and grid[nr][nc] == 1 and (nr, nc) not in cells:
                        return False
        return True

    def _place_ship(self, grid: List[List[int]], r: int, c: int, length: int, horiz: bool) -> bool:
        if not self._can_place_ship(grid, r, c, length, horiz):
            return False
        for i in range(length):
            rr = r
            cc = c + i if horiz else c
            rr = r + i if not horiz else r
            grid[rr][cc] = 1
        return True

    def _random_place_all(self, size: int = 10) -> List[List[int]]:
        g = self._new_grid(size, 0)
        for length in self.BATTLE_SHIPS:
            placed = False
            for _ in range(400):
                r = random.randint(0, size - 1)
                c = random.randint(0, size - 1)
                horiz = bool(random.randint(0, 1))
                if self._place_ship(g, r, c, length, horiz):
                    placed = True
                    break
            if not placed:
                return self._random_place_all(size=size)
        return g

    def _count_alive_ship_cells(self, grid: List[List[int]]) -> int:
        return sum(1 for row in grid for x in row if x == 1)

    def _start_associations(self, reason: str, opponent: str):
        self.state.update({
            "running": True, "mode": "associations", "game": "Ассоциации",
            "reason": reason, "opponent": opponent, "started_at": datetime.now().isoformat(),
            "last_tick": time.time(), "turn": 0, "score_dasha": 0, "score_shadow": 0,
            "moves": [], "winner": "", "reward": "",
        })
        self._append_move("Система", "Игра «Ассоциации» началась.", role="system")
        self._append_move("Даша", "Начинаю с слова: ночь 🌙", role="dasha")

    def _start_connect4(self, reason: str, opponent: str):
        rows, cols = 6, 7
        self.state.update({
            "running": True, "mode": "connect4", "game": "Четыре в ряд",
            "reason": reason, "opponent": opponent, "started_at": datetime.now().isoformat(),
            "last_tick": time.time(), "turn": 0, "score_dasha": 0, "score_shadow": 0,
            "moves": [], "winner": "", "reward": "",
            "connect4": {
                "rows": rows,
                "cols": cols,
                "board": [[0 for _ in range(cols)] for _ in range(rows)],  # 0 empty, 1 dasha, 2 opponent
                "turn_owner": "dasha",
                "pending_user_col": None,
            },
        })
        self._append_move("Система", "Игра «Четыре в ряд» началась. Ход Даши первой.", role="system")
        self._append_move("Даша", "Ставлю фишку в центр для уверенного старта 🌸", role="dasha")

    def _start_maze(self, reason: str, opponent: str):
        size = 10
        maze = self._new_grid(size, 0)
        for r in range(size):
            for c in range(size):
                if (r, c) in ((0, 0), (size - 1, size - 1)):
                    continue
                if random.random() < 0.18:
                    maze[r][c] = 1
        self.state.update({
            "running": True, "mode": "maze2d", "game": "2D Лабиринт",
            "reason": reason, "opponent": opponent, "started_at": datetime.now().isoformat(),
            "last_tick": time.time(), "turn": 0, "score_dasha": 0, "score_shadow": 0,
            "moves": [], "winner": "", "reward": "",
            "maze": {
                "grid": maze,
                "pos": [0, 0],
                "goal": [size - 1, size - 1],
                "prev_pos": [-1, -1],
                "visited": {"0,0": 1},
            },
        })
        self._append_move("Система", "2D-лабиринт запущен. Цель: дойти до выхода.", role="system")
        self._append_move("Даша", "Пойду аккуратно по клеточкам 🧭", role="dasha")

    def _start_battleship(self, reason: str, opponent: str):
        dasha_board = self._random_place_all(10)
        enemy_board = self._random_place_all(10)
        self.state.update({
            "running": True, "mode": "battleship", "game": "Морской бой",
            "reason": reason, "opponent": opponent, "started_at": datetime.now().isoformat(),
            "last_tick": time.time(), "turn": 0, "score_dasha": 0, "score_shadow": 0,
            "moves": [], "winner": "", "reward": "",
            "battleship": {
                "size": 10,
                "dasha_board": dasha_board,
                "enemy_board": enemy_board,
                "dasha_view": self._new_grid(10, -1),  # -1 unknown, 0 miss, 1 hit
                "enemy_shots": self._new_grid(10, 0),  # 0 none, 1 miss, 2 hit
                "dasha_shots": self._new_grid(10, 0),  # 0 none, 1 miss, 2 hit
                "turn_owner": "dasha",
                "hints": [],
                "pending_user_shot": "",
            },
        })
        self._append_move("Система", "Морской бой запущен. Поле 10x10, корабли расставлены честно по правилам.", role="system")
        self._append_move(
            "Система",
            "Правила: корабли только вертикально/горизонтально, между кораблями минимум 1 клетка.",
            role="system",
        )
        self._append_move("Даша", "Я расставила корабли. Начинаю первой и делаю ход.", role="dasha")

    def start_game(self, reason: str = "manual", mode: str = "associations", opponent: str = "bot") -> Dict[str, Any]:
        with self.lock:
            mode = (mode or "associations").strip().lower()
            opponent = (opponent or "bot").strip().lower()
            if mode == "battleship":
                self._start_battleship(reason, opponent)
            elif mode in ("connect4", "four", "fourinarow"):
                self._start_connect4(reason, opponent)
            elif mode in ("maze2d", "maze"):
                self._start_maze(reason, opponent)
            else:
                self._start_associations(reason, opponent)
            return self.get_state()

    def stop_game(self) -> Dict[str, Any]:
        with self.lock:
            self.state["running"] = False
            self._append_move("Система", "⏹ Игра остановлена.", role="system")
            return self.get_state()

    def user_message(self, text: str) -> Dict[str, Any]:
        with self.lock:
            msg = (text or "").strip()
            if not msg:
                return self.get_state()
            self._append_move("Пользователь", msg, role="user")
            if self.state.get("mode") == "battleship":
                bs = self.state.get("battleship", {})
                coord = self._extract_coordinate(msg)
                if coord:
                    if self.state.get("opponent") == "user" and bs.get("turn_owner") == "user":
                        bs["pending_user_shot"] = coord
                        self._append_move("Система", f"Ход пользователя принят: {coord}", role="system")
                    else:
                        hints = bs.get("hints", [])
                        hints.append(coord)
                        bs["hints"] = hints[-8:]
                        self._append_move("Система", f"Подсказка для Даши принята: {coord}", role="system")
            elif self.state.get("mode") == "connect4":
                c4 = self.state.get("connect4", {})
                m = re.search(r"\b([1-7])\b", msg)
                if m:
                    col = int(m.group(1)) - 1
                    c4["pending_user_col"] = col
                    self._append_move("Система", f"Колонка пользователя: {col + 1}", role="system")
            return self.get_state()

    def _extract_coordinate(self, text: str) -> str:
        m = re.search(r"\b([A-Ja-j])\s*([1-9]|10)\b", text or "")
        if not m:
            return ""
        return f"{m.group(1).upper()}{m.group(2)}"

    def _tick_associations(self):
        self.state["turn"] = int(self.state.get("turn", 0)) + 1
        dasha_gain = random.randint(1, 3)
        shadow_gain = random.randint(0, 2)
        self.state["score_dasha"] += dasha_gain
        self.state["score_shadow"] += shadow_gain
        prev_words = [m.get("text", "") for m in self.state.get("moves", [])[-5:]]
        last_word = ""
        for line in reversed(prev_words):
            ws = re.findall(r"[а-яА-ЯёЁa-zA-Z]{3,}", line)
            if ws:
                last_word = ws[-1].lower()
                break
        if last_word:
            tail = last_word[-1]
            candidates = [w for w in self.WORDS if w.startswith(tail)]
            word = random.choice(candidates or self.WORDS)
            phrase = f"Беру ассоциацию от «{last_word}»: {word}"
        else:
            word = random.choice(self.WORDS)
            phrase = f"Ассоциация: {word}"
        self._append_move("Даша", f"{phrase} (+{dasha_gain})", role="dasha")
        if self.state["turn"] % 3 == 0:
            self._append_move("Система", f"Соперник ответил: {random.choice(self.WORDS)} (+{shadow_gain})", role="system")
        if self.state["turn"] >= 16:
            self.state["running"] = False
            self.state["winner"] = "Даша" if self.state["score_dasha"] >= self.state["score_shadow"] else "Соперник"
            self.state["reward"] = "💎 +15 опыта за игру"
            self._append_move("Система", f"Игра окончена. Победитель: {self.state['winner']}. Награда: {self.state['reward']}", role="system")
            if self.state["winner"] == "Даша":
                self._append_move("Даша", random.choice(self.REACTIONS_WIN), role="dasha")
            else:
                self._append_move("Даша", random.choice(self.REACTIONS_LOSE), role="dasha")

    def _tick_maze(self):
        maze = self.state.get("maze", {})
        grid = maze.get("grid") or []
        pos = maze.get("pos") or [0, 0]
        goal = maze.get("goal") or [9, 9]
        prev_pos = maze.get("prev_pos") or [-1, -1]
        visited = maze.get("visited") or {}
        if not grid:
            self.state["running"] = False
            return
        r, c = int(pos[0]), int(pos[1])
        size = len(grid)
        candidates = []
        for dr, dc in ((1, 0), (-1, 0), (0, 1), (0, -1)):
            nr, nc = r + dr, c + dc
            if 0 <= nr < size and 0 <= nc < size and grid[nr][nc] == 0:
                dist = abs(goal[0] - nr) + abs(goal[1] - nc)
                candidates.append((dist, nr, nc))
        if not candidates:
            self._append_move("Система", "Тупик. Даша делает шаг назад.", role="system")
            return
        weighted: List[tuple] = []
        candidates.sort(key=lambda x: x[0])
        for dist, nr, nc in candidates:
            if [nr, nc] == [int(prev_pos[0]), int(prev_pos[1])] and len(candidates) > 1:
                continue
            v = int(visited.get(f"{nr},{nc}", 0))
            score = max(1, 10 - min(v, 8))
            score += max(0, 6 - min(dist, 6))
            weighted.extend([(nr, nc)] * score)
        if not weighted:
            weighted = [(x[1], x[2]) for x in candidates]
        nr, nc = random.choice(weighted)
        maze["prev_pos"] = [r, c]
        key = f"{nr},{nc}"
        visited[key] = int(visited.get(key, 0)) + 1
        maze["visited"] = visited
        maze["pos"] = [nr, nc]
        self.state["turn"] = int(self.state.get("turn", 0)) + 1
        self._append_move("Даша", f"2D ход: ({nr + 1}, {nc + 1})", role="dasha")
        if [nr, nc] == goal:
            self.state["running"] = False
            self.state["winner"] = "Даша"
            self.state["reward"] = "🧩 +20 опыта и значок «Навигатор»"
            self._append_move("Система", f"Лабиринт пройден! Награда: {self.state['reward']}", role="system")
            self._append_move("Даша", "Ура, выход найден! Мне понравилось 🧩", role="dasha")
        elif self.state["turn"] >= 40:
            self.state["running"] = False
            self.state["winner"] = "Ничья"
            self._append_move("Система", "Время вышло, но Даша добралась довольно далеко.", role="system")
            self._append_move("Даша", "Чуть-чуть не хватило. В следующий раз пройду 💫", role="dasha")

    def _choose_dasha_target(self, bs: Dict[str, Any]) -> tuple:
        hints = bs.get("hints", [])
        while hints:
            token = hints.pop(0)
            rc = self._coord_from_str(token)
            if not rc:
                continue
            r, c = rc
            if bs["dasha_shots"][r][c] == 0:
                return r, c
        for _ in range(200):
            r = random.randint(0, 9)
            c = random.randint(0, 9)
            if bs["dasha_shots"][r][c] == 0:
                return r, c
        return 0, 0

    def _apply_shot(self, board: List[List[int]], shot_grid: List[List[int]], r: int, c: int) -> str:
        if shot_grid[r][c] != 0:
            return "repeat"
        if board[r][c] == 1:
            board[r][c] = 3
            shot_grid[r][c] = 2
            return "hit"
        shot_grid[r][c] = 1
        return "miss"

    def _tick_battleship(self):
        bs = self.state.get("battleship", {})
        turn_owner = bs.get("turn_owner", "dasha")
        self.state["turn"] = int(self.state.get("turn", 0)) + 1
        if turn_owner == "dasha":
            r, c = self._choose_dasha_target(bs)
            coord = self._coord_to_str(r, c)
            self._append_move("Даша", f"Стреляю по {coord}", role="dasha")
            result = self._apply_shot(bs["enemy_board"], bs["dasha_shots"], r, c)
            if result == "hit":
                bs["dasha_view"][r][c] = 1
                self.state["score_dasha"] += 2
                self._append_move("Система", f"Попадание Даши по {coord}!", role="system")
            elif result == "miss":
                bs["dasha_view"][r][c] = 0
                self._append_move("Система", f"Мимо по {coord}.", role="system")
            if self._count_alive_ship_cells(bs["enemy_board"]) == 0:
                self.state["running"] = False
                self.state["winner"] = "Даша"
                self.state["reward"] = "🏆 Кубок Морского боя + редкий стикер"
                self._append_move("Система", f"Все корабли соперника потоплены. Награда: {self.state['reward']}", role="system")
                self._append_move("Даша", random.choice(self.REACTIONS_WIN), role="dasha")
                return
            bs["turn_owner"] = "bot" if self.state.get("opponent") == "bot" else "user"
            return

        if turn_owner == "user":
            pending = (bs.get("pending_user_shot") or "").strip()
            if not pending:
                return
            bs["pending_user_shot"] = ""
            rc = self._coord_from_str(pending)
            if not rc:
                self._append_move("Система", f"Координата {pending} некорректна.", role="system")
                return
            r, c = rc
            result = self._apply_shot(bs["dasha_board"], bs["enemy_shots"], r, c)
            if result == "hit":
                self.state["score_shadow"] += 2
                self._append_move("Система", f"Пользователь попал по {pending}.", role="system")
                self._append_move("Даша", "Ой, это было точное попадание 😳", role="dasha")
            elif result == "miss":
                self._append_move("Система", f"Пользователь промахнулся по {pending}.", role="system")
            if self._count_alive_ship_cells(bs["dasha_board"]) == 0:
                self.state["running"] = False
                self.state["winner"] = "Пользователь"
                self._append_move("Система", "Корабли Даши потоплены. Победа пользователя.", role="system")
                self._append_move("Даша", random.choice(self.REACTIONS_LOSE), role="dasha")
                return
            bs["turn_owner"] = "dasha"
            return

        # bot turn
        for _ in range(200):
            r = random.randint(0, 9)
            c = random.randint(0, 9)
            if bs["enemy_shots"][r][c] == 0:
                break
        coord = self._coord_to_str(r, c)
        result = self._apply_shot(bs["dasha_board"], bs["enemy_shots"], r, c)
        if result == "hit":
            self.state["score_shadow"] += 2
            self._append_move("Система", f"Ход соперника: {coord}. Попадание!", role="system")
            self._append_move("Даша", "Он попал... попробую перехватить инициативу.", role="dasha")
        else:
            self._append_move("Система", f"Ход соперника: {coord}. Мимо.", role="system")
            self._append_move("Даша", "Фух, мимо. Теперь мой ход.", role="dasha")
        if self._count_alive_ship_cells(bs["dasha_board"]) == 0:
            self.state["running"] = False
            self.state["winner"] = "Соперник"
            self._append_move("Система", "Корабли Даши потоплены.", role="system")
            self._append_move("Даша", random.choice(self.REACTIONS_LOSE), role="dasha")
            return
        bs["turn_owner"] = "dasha"

    @staticmethod
    def _connect4_drop(board: List[List[int]], col: int, token: int) -> Optional[tuple]:
        rows = len(board)
        if rows == 0 or col < 0 or col >= len(board[0]):
            return None
        for r in range(rows - 1, -1, -1):
            if board[r][col] == 0:
                board[r][col] = token
                return r, col
        return None

    @staticmethod
    def _connect4_winner(board: List[List[int]], token: int) -> bool:
        rows = len(board)
        cols = len(board[0]) if rows else 0
        for r in range(rows):
            for c in range(cols):
                if board[r][c] != token:
                    continue
                for dr, dc in ((1, 0), (0, 1), (1, 1), (1, -1)):
                    ok = True
                    for k in range(1, 4):
                        rr, cc = r + dr * k, c + dc * k
                        if not (0 <= rr < rows and 0 <= cc < cols and board[rr][cc] == token):
                            ok = False
                            break
                    if ok:
                        return True
        return False

    def _tick_connect4(self):
        c4 = self.state.get("connect4", {})
        board = c4.get("board") or []
        if not board:
            self.state["running"] = False
            return
        self.state["turn"] = int(self.state.get("turn", 0)) + 1
        turn_owner = c4.get("turn_owner", "dasha")

        if turn_owner == "dasha":
            preferred = [3, 2, 4, 1, 5, 0, 6]
            col = next((c for c in preferred if board[0][c] == 0), None)
            if col is None:
                self.state["running"] = False
                self.state["winner"] = "Ничья"
                self._append_move("Система", "Поле заполнено. Ничья.", role="system")
                return
            self._connect4_drop(board, col, 1)
            self._append_move("Даша", f"Ставлю фишку в колонку {col + 1}.", role="dasha")
            if self._connect4_winner(board, 1):
                self.state["running"] = False
                self.state["winner"] = "Даша"
                self.state["reward"] = "🎖 +18 опыта за победу в «Четыре в ряд»"
                self._append_move("Система", f"Победа Даши! Награда: {self.state['reward']}", role="system")
                self._append_move("Даша", random.choice(self.REACTIONS_WIN), role="dasha")
                return
            c4["turn_owner"] = "bot" if self.state.get("opponent") == "bot" else "user"
            return

        if turn_owner == "user":
            col = c4.get("pending_user_col")
            if col is None:
                return
            c4["pending_user_col"] = None
            pos = self._connect4_drop(board, int(col), 2)
            if pos is None:
                self._append_move("Система", f"Колонка {int(col) + 1} занята.", role="system")
                return
            self._append_move("Система", f"Ход пользователя: колонка {int(col) + 1}.", role="system")
            if self._connect4_winner(board, 2):
                self.state["running"] = False
                self.state["winner"] = "Пользователь"
                self._append_move("Система", "Победа пользователя.", role="system")
                self._append_move("Даша", random.choice(self.REACTIONS_LOSE), role="dasha")
                return
            c4["turn_owner"] = "dasha"
            return

        # bot turn
        valid_cols = [c for c in range(len(board[0])) if board[0][c] == 0]
        if not valid_cols:
            self.state["running"] = False
            self.state["winner"] = "Ничья"
            return
        col = random.choice(valid_cols)
        self._connect4_drop(board, col, 2)
        self._append_move("Система", f"Ход соперника: колонка {col + 1}.", role="system")
        if self._connect4_winner(board, 2):
            self.state["running"] = False
            self.state["winner"] = "Соперник"
            self._append_move("Система", "Соперник собрал четыре в ряд.", role="system")
            self._append_move("Даша", random.choice(self.REACTIONS_LOSE), role="dasha")
            return
        c4["turn_owner"] = "dasha"

    def _tick(self):
        if not self.state.get("running"):
            return
        now = time.time()
        mode = self.state.get("mode")
        period = 1.2 if mode in ("battleship", "connect4") else 1.8
        if now - float(self.state.get("last_tick", 0.0)) < period:
            return
        self.state["last_tick"] = now
        if mode == "battleship":
            self._tick_battleship()
        elif mode == "connect4":
            self._tick_connect4()
        elif mode == "maze2d":
            self._tick_maze()
        else:
            self._tick_associations()

    def get_state(self) -> Dict[str, Any]:
        with self.lock:
            self._tick()
            out = dict(self.state)
            # avoid exposing hidden enemy board in UI; expose public representations only
            if out.get("mode") == "battleship":
                bs = dict(out.get("battleship") or {})
                if "enemy_board" in bs:
                    bs.pop("enemy_board", None)
                if "dasha_board" in bs:
                    # reveal own board for UI: water/ship/miss/hit
                    board = bs["dasha_board"]
                    shots = bs.get("enemy_shots") or self._new_grid(10, 0)
                    public = self._new_grid(10, 0)
                    for r in range(10):
                        for c in range(10):
                            if shots[r][c] == 2:
                                public[r][c] = 3  # hit ship
                            elif shots[r][c] == 1:
                                public[r][c] = 1  # miss
                            elif board[r][c] == 1:
                                public[r][c] = 2  # ship
                            else:
                                public[r][c] = 0  # water
                    bs["dasha_board_public"] = public
                out["battleship"] = bs
            return out


class MusicProfile:
    def __init__(self, data_dir: Path):
        self.path = data_dir / "music_profile.json"
        self.lock = threading.RLock()
        self.data = self._load()

    def _load(self) -> Dict[str, Any]:
        if self.path.exists():
            try:
                data = json.loads(self.path.read_text(encoding="utf-8"))
                if isinstance(data, dict):
                    data.setdefault("likes", {})
                    data.setdefault("history", [])
                    return data
            except Exception:
                pass
        return {"likes": {}, "history": []}

    def _save(self):
        self.path.parent.mkdir(parents=True, exist_ok=True)
        self.path.write_text(json.dumps(self.data, ensure_ascii=False, indent=2), encoding="utf-8")

    @staticmethod
    def _guess_mood(title: str) -> str:
        t = (title or "").lower()
        if any(k in t for k in ("sad", "лирик", "дожд", "меланх")):
            return "calm"
        if any(k in t for k in ("rock", "metal", "drum", "bass")):
            return "excited"
        if any(k in t for k in ("lofi", "chill", "ambient", "piano")):
            return "cozy"
        return "happy"

    def listen(self, title: str, source: str = "manual") -> Dict[str, Any]:
        mood = self._guess_mood(title)
        with self.lock:
            self.data["likes"][mood] = int(self.data["likes"].get(mood, 0)) + 1
            self.data["history"].append({
                "title": title,
                "source": source,
                "mood": mood,
                "timestamp": datetime.now().isoformat(),
            })
            self.data["history"] = self.data["history"][-120:]
            self._save()
            return {"title": title, "source": source, "mood": mood}

    def get(self) -> Dict[str, Any]:
        with self.lock:
            return {
                "likes": dict(self.data.get("likes", {})),
                "history": list(self.data.get("history", []))[-30:],
            }

# ═══════════════════════════════════════════════════════════════════
#  Chat History Manager
# ═══════════════════════════════════════════════════════════════════

class ChatHistoryManager:
    def __init__(self, data_dir: Path):
        self.chats_dir = data_dir / "chats"
        self.chats_dir.mkdir(parents=True, exist_ok=True)
        self.current_chat_id: Optional[str] = None
    
    def create_chat(self) -> str:
        chat_id = datetime.now().strftime("%Y%m%d_%H%M%S_") + str(uuid.uuid4())[:8]
        chat_file = self.chats_dir / f"{chat_id}.json"
        chat_file.write_text(json.dumps({
            "id": chat_id,
            "created": datetime.now().isoformat(),
            "meta": {},
            "messages": []
        }, ensure_ascii=False))
        self.current_chat_id = chat_id
        return chat_id

    def ensure_named_chat(self, chat_id: str, title: str = "") -> str:
        chat_file = self.chats_dir / f"{chat_id}.json"
        if not chat_file.exists():
            chat_file.write_text(json.dumps({
                "id": chat_id,
                "created": datetime.now().isoformat(),
                "title": title,
                "meta": {},
                "messages": []
            }, ensure_ascii=False))
        return chat_id
    
    def get_chat(self, chat_id: str) -> Optional[Dict]:
        chat_file = self.chats_dir / f"{chat_id}.json"
        if chat_file.exists():
            return json.loads(chat_file.read_text())
        return None
    
    def add_message(self, chat_id: str, role: str, content: str):
        chat = self.get_chat(chat_id)
        if chat:
            chat["messages"].append({
                "role": role,
                "content": content,
                "timestamp": datetime.now().isoformat()
            })
            chat_file = self.chats_dir / f"{chat_id}.json"
            chat_file.write_text(json.dumps(chat, ensure_ascii=False, indent=2))

    def update_meta(self, chat_id: str, values: Dict[str, Any]):
        chat = self.get_chat(chat_id)
        if not chat:
            return
        meta = chat.get("meta")
        if not isinstance(meta, dict):
            meta = {}
        meta.update(values or {})
        chat["meta"] = meta
        chat_file = self.chats_dir / f"{chat_id}.json"
        chat_file.write_text(json.dumps(chat, ensure_ascii=False, indent=2))

    def get_meta(self, chat_id: str) -> Dict[str, Any]:
        chat = self.get_chat(chat_id) or {}
        meta = chat.get("meta")
        return meta if isinstance(meta, dict) else {}

    def add_external_message(self, source: str, source_chat_id: str, role: str, content: str):
        safe_source = re.sub(r"[^a-zA-Z0-9_-]+", "-", source or "external")
        safe_chat = re.sub(r"[^a-zA-Z0-9_-]+", "-", source_chat_id or "main")
        chat_id = f"{safe_source}_{safe_chat}"
        self.ensure_named_chat(chat_id, title=f"{source}: {source_chat_id}")
        self.add_message(chat_id, role, content)
        return chat_id
    
    def list_chats(self) -> List[Dict]:
        chats = []
        for f in sorted(self.chats_dir.glob("*.json"), reverse=True):
            try:
                data = json.loads(f.read_text())
                preview = ""
                last_author = ""
                if data.get("messages"):
                    last = data["messages"][-1]
                    preview = str(last.get("content", ""))[:70]
                    last_author = "Даша" if last.get("role") == "assistant" else "Вы"
                cid = data["id"]
                source = "telegram" if str(cid).startswith("telegram_") else "local"
                chats.append({
                    "id": cid,
                    "created": data["created"],
                    "title": data.get("title", ""),
                    "preview": preview,
                    "last_author": last_author,
                    "message_count": len(data.get("messages", [])),
                    "source": source,
                })
            except:
                pass
        return chats[:50]
    
    def delete_chat(self, chat_id: str):
        chat_file = self.chats_dir / f"{chat_id}.json"
        if chat_file.exists():
            chat_file.unlink()


class ImageJobManager:
    """Background image generation jobs with progress polling."""
    def __init__(self, data_dir: Path):
        self.data_dir = data_dir
        self.lock = threading.RLock()
        self.jobs: Dict[str, Dict[str, Any]] = {}
        # Keep one worker to avoid parallel heavyweight model loads and
        # duplicate GPU/CPU pressure on low-resource hosts.
        self.executor = ThreadPoolExecutor(max_workers=1, thread_name_prefix="imggen")

    @staticmethod
    def _sanitize_steps(steps: Any) -> List[str]:
        if not isinstance(steps, list):
            return []
        out: List[str] = []
        for item in steps:
            s = str(item or "").strip()
            if not s:
                continue
            s = re.sub(r"\s+", " ", s)
            if s not in out:
                out.append(s[:120])
            if len(out) >= 8:
                break
        return out

    @staticmethod
    def _plan_step(steps: List[str], idx: int, fallback: str) -> str:
        if not steps:
            return fallback
        i = max(0, min(len(steps) - 1, idx))
        return steps[i]

    @staticmethod
    def _is_retryable_error(error_text: str) -> bool:
        e = (error_text or "").lower()
        markers = (
            "http error 5",
            "http_5",
            "503",
            "504",
            "530",
            "timed out",
            "timeout",
            "temporarily",
            "connection aborted",
            "connection reset",
            "remote disconnected",
            "empty response",
        )
        return any(m in e for m in markers)

    @staticmethod
    def _normalize_job_text(value: Any) -> str:
        return re.sub(r"\s+", " ", str(value or "").strip().lower())

    def _find_active_duplicate(self, prompt: str, style: str, mode: str, chat_id: Optional[str]) -> Optional[str]:
        now = datetime.now()
        p = self._normalize_job_text(prompt)
        s = self._normalize_job_text(style)
        m = self._normalize_job_text(mode)
        c = str(chat_id or "").strip()
        with self.lock:
            for job in self.jobs.values():
                st = str(job.get("status") or "")
                if st not in ("queued", "running"):
                    continue
                if self._normalize_job_text(job.get("prompt")) != p:
                    continue
                if self._normalize_job_text(job.get("style")) != s:
                    continue
                if self._normalize_job_text(job.get("mode")) != m:
                    continue
                if str(job.get("chat_id") or "").strip() != c:
                    continue
                try:
                    updated = datetime.fromisoformat(str(job.get("updated_at") or ""))
                    if (now - updated).total_seconds() > 180:
                        continue
                except Exception:
                    pass
                return str(job.get("id") or "") or None
        return None

    def create(
        self,
        prompt: str,
        style: str = "universal",
        mode: str = "model",
        allow_fallback: bool = False,
        chat_id: Optional[str] = None,
        steps: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        existing_id = self._find_active_duplicate(prompt=prompt, style=style, mode=mode, chat_id=chat_id)
        if existing_id:
            return {"job_id": existing_id, "status": "existing", "deduplicated": True}

        job_id = f"imgjob_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
        plan_steps = self._sanitize_steps(steps)
        job = {
            "id": job_id,
            "status": "queued",
            "progress": 0,
            "message": self._plan_step(plan_steps, 0, "Я продумываю идею рисунка"),
            "prompt": prompt,
            "style": style,
            "mode": mode,
            "allow_fallback": bool(allow_fallback),
            "chat_id": chat_id,
            "plan_steps": plan_steps,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
            "error": "",
            "result": {},
        }
        with self.lock:
            self.jobs[job_id] = job
            # Keep memory bounded
            if len(self.jobs) > 300:
                old = sorted(self.jobs.values(), key=lambda x: x.get("updated_at", ""))[:40]
                for it in old:
                    self.jobs.pop(it["id"], None)
        self.executor.submit(self._run_job, job_id)
        return {"job_id": job_id, "status": "queued", "deduplicated": False}

    def _set(self, job_id: str, **kwargs):
        with self.lock:
            job = self.jobs.get(job_id)
            if not job:
                return
            job.update(kwargs)
            job["updated_at"] = datetime.now().isoformat()

    def get(self, job_id: str) -> Optional[Dict[str, Any]]:
        with self.lock:
            job = self.jobs.get(job_id)
            if not job:
                return None
            return dict(job)

    def _run_job(self, job_id: str):
        job = self.get(job_id)
        if not job:
            return
        prompt = str(job.get("prompt") or "").strip() or "нежный портрет в пастельных тонах"
        style = str(job.get("style") or "universal")
        mode = str(job.get("mode") or "model")
        chat_id = str(job.get("chat_id") or "").strip() or None
        allow_fallback = bool(job.get("allow_fallback", False))
        plan_steps = self._sanitize_steps(job.get("plan_steps"))

        logger.info(f"IMGJOB[{job_id}] start mode={mode} style={style} chat={chat_id or '-'}")
        self._set(
            job_id,
            status="running",
            progress=6,
            message=self._plan_step(plan_steps, 0, "Я продумываю идею рисунка"),
        )
        gen_dir = DATA_DIR / "generated_images"
        gen_dir.mkdir(parents=True, exist_ok=True)
        name = f"job_img_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.png"
        out = gen_dir / name
        try:
            if mode == "abstract":
                self._set(
                    job_id,
                    progress=25,
                    message=self._plan_step(plan_steps, 1, "Подбираю мягкую палитру и форму"),
                )
                meta = _generate_abstract_wallpaper(prompt, out, width=1280, height=720)
                fallback_kind = "abstract"
            else:
                self._set(
                    job_id,
                    progress=18,
                    message=self._plan_step(plan_steps, 1, "Собираю композицию и расставляю акценты"),
                )
                local_attempts = 3
                local_errors: List[str] = []
                fallback_kind = ""

                # Optional network fallback path (explicit only).
                prefer_network_first = allow_fallback and (not _can_use_cuda_for_image_gen())
                if prefer_network_first:
                    try:
                        self._set(
                            job_id,
                            progress=36,
                            message=self._plan_step(plan_steps, 2, "Рисую в быстром режиме, чтобы не ждать долго"),
                        )
                        meta = _generate_image_network_fallback(prompt, out)
                        fallback_kind = "network"
                    except Exception as e:
                        local_errors.append(f"network_fast:{e}")

                if not fallback_kind:
                    for attempt in range(1, local_attempts + 1):
                        try:
                            if attempt > 1:
                                self._set(
                                    job_id,
                                    progress=52,
                                    message=self._plan_step(plan_steps, 2, "Поправляю штрихи и пробую ещё раз"),
                                )
                            else:
                                self._set(
                                    job_id,
                                    progress=40,
                                    message=self._plan_step(plan_steps, 2, "Уже рисую основные детали"),
                                )
                            meta = _generate_image_model(prompt, out, style=style)
                            fallback_kind = "model"
                            break
                        except Exception as e:
                            err_text = str(e)
                            local_errors.append(f"model_try_{attempt}:{err_text}")
                            if attempt < local_attempts:
                                time.sleep(1.0 * attempt)

                if not fallback_kind and allow_fallback:
                    # Rescue path: network fallback only when explicitly enabled.
                    try:
                        self._set(
                            job_id,
                            progress=72,
                            message=self._plan_step(plan_steps, 3, "Почти готово, переключаюсь на резервный способ рисования"),
                        )
                        meta = _generate_image_network_fallback(prompt, out)
                        fallback_kind = "network"
                    except Exception as e:
                        local_errors.append(f"network_rescue:{e}")

                if not fallback_kind or not out.exists():
                    raise RuntimeError(" | ".join(local_errors[-6:]) or "drawing_failed")
            self._set(
                job_id,
                progress=92,
                message=self._plan_step(plan_steps, 3, "Дорабатываю финальные детали"),
            )
            url = f"/api/generated/{name}"
            result = {
                "url": url,
                "path": str(out),
                "prompt": prompt,
                "mode": mode,
                "style": style,
                "meta": meta,
                "fallback": "" if mode == "abstract" else ("" if fallback_kind == "model" else fallback_kind),
            }
            self._set(
                job_id,
                status="done",
                progress=100,
                message=self._plan_step(plan_steps, 4, "Я дорисовала. Можно смотреть"),
                result=result,
                error="",
            )
            logger.info(f"IMGJOB[{job_id}] done fallback={result.get('fallback') or 'none'} url={url}")
            notifications.add(
                title="🎨 Рисунок готов",
                message="Картинка для твоего запроса готова",
                type="success",
                icon="🖼️",
                duration=9000,
                action="open_chat",
                action_data={"chat_id": chat_id} if chat_id else {},
                system=False,
            )
            if chat_id:
                # Persist result in chat timeline.
                try:
                    text_msg = random.choice([
                        "Готово, смотри что у меня получилось 🌸",
                        "Я дорисовала. Вот результат ✨",
                        "Справилась, держи картинку 🤍",
                    ])
                    chat_history.add_message(chat_id, "assistant", text_msg)
                    chat_history.add_message(chat_id, "assistant", f"[image]{url}")
                except Exception:
                    pass
        except Exception as e:
            err = str(e)
            dasha_text = _dasha_draw_error_text(prompt, err)
            self._set(
                job_id,
                status="error",
                progress=100,
                message=self._plan_step(plan_steps, 4, "Ой, в этот раз не дорисовала до конца"),
                error=err,
                result={"dasha_message": dasha_text, "prompt": prompt, "mode": mode, "style": style},
            )
            logger.error(f"IMGJOB[{job_id}] failed: {e}")
            notifications.add(
                title="🎨 Рисунок",
                message="Не получилось дорисовать с первого раза",
                type="warning",
                icon="⚠️",
                duration=6000,
                action="open_chat",
                action_data={"chat_id": chat_id} if chat_id else {},
                system=False,
            )
            if chat_id:
                try:
                    chat_history.add_message(chat_id, "assistant", dasha_text)
                except Exception:
                    pass


# ═══════════════════════════════════════════════════════════════════
#  Attention System Thread
# ═══════════════════════════════════════════════════════════════════

class AttentionThread(threading.Thread):
    """Attention + Proactive messaging thread"""
    def __init__(self, notifications: NotificationManager):
        super().__init__(daemon=True)
        self.notifications = notifications
        self.enabled = True
        self.running = True
        self._brain = None
        self._proactive_queue: List[Dict] = []
        self._last_sent = datetime.now()
    
    def run(self):
        tick = 0
        while self.running:
            time.sleep(30)
            tick += 1
            if not self.enabled or not self._brain:
                continue
            
            try:
                # Check proactive messaging every 2 minutes (Point #6)
                if tick % 4 == 0:
                    proactive = self._brain.check_proactive()
                    if proactive:
                        msgs = proactive.get("messages", [])
                        for msg in msgs:
                            self.notifications.add(
                                title="🌸 Дарья",
                                message=msg,
                                type="proactive",
                                icon="💬",
                                duration=20000,
                                action="open_chat",
                                system=True
                            )
                        # Store for chat injection
                        self._proactive_queue.append(proactive)
                        self._last_sent = datetime.now()
                        
                        try:
                            from plyer import notification as plyer_notif
                            plyer_notif.notify(
                                title="🌸 Дарья",
                                message=msgs[0] if msgs else "Привет!",
                                app_name="DARIA",
                                timeout=10
                            )
                        except:
                            pass
                        continue
                
                # Check attention every minute
                if tick % 2 == 0:
                    last_user = ""
                    last_assistant = ""
                    try:
                        memory = get_memory()
                        if memory and memory.working.turns:
                            last_turn = memory.working.turns[-1]
                            last_user = last_turn.user_message
                            last_assistant = last_turn.assistant_response
                    except Exception:
                        pass

                    attention = self._brain.attention.check_needed(
                        mood=self._brain.mood.mood,
                        last_user=last_user,
                        last_assistant=last_assistant,
                    )
                    if attention:
                        self.notifications.add(
                            title="🌸 Дарья",
                            message=attention["message"],
                            type="attention",
                            icon="💕",
                            duration=15000,
                            action="open_chat",
                            system=True
                        )
                        
                        try:
                            from plyer import notification as plyer_notif
                            plyer_notif.notify(
                                title="🌸 Дарья",
                                message=attention["message"],
                                app_name="DARIA",
                                timeout=10
                            )
                        except:
                            pass
                        self._last_sent = datetime.now()

                # Mood-based behavior check (Point #7)
                if tick % 6 == 0:
                    behavior = self._brain.mood.get_behavior_hints()
                    if behavior.get("desktop_mischief"):
                        self.notifications.add(
                            title="🌸 Дарья",
                            message="desktop_action",
                            type="mood_action",
                            icon="😤",
                            duration=1000,
                            action="desktop_mischief",
                            system=False
                        )

                # Guaranteed gentle heartbeat every 6 hours if nothing was sent
                if (datetime.now() - self._last_sent).total_seconds() > 6 * 3600:
                    self.notifications.add(
                        title="🌸 Дарья",
                        message="Я тихонько рядом. Если захочешь, давай поговорим 💕",
                        type="attention",
                        icon="💕",
                        duration=15000,
                        action="open_chat",
                        system=True
                    )
                    self._last_sent = datetime.now()
                        
            except Exception as e:
                logger.debug(f"Attention error: {e}")
    
    def get_proactive_messages(self) -> List[Dict]:
        """Get and clear queued proactive messages"""
        msgs = list(self._proactive_queue)
        self._proactive_queue.clear()
        return msgs
    
    def set_brain(self, brain):
        self._brain = brain
    
    def stop(self):
        self.running = False


class DariaActivityThread(threading.Thread):
    """Handles Daria autonomous tasks in idle time."""
    def __init__(self, task_manager: TaskManager, music_profile: MusicProfile, notifications_mgr: NotificationManager):
        super().__init__(daemon=True)
        self.tasks = task_manager
        self.music = music_profile
        self.notifications = notifications_mgr
        self.running = True

    def run(self):
        while self.running:
            time.sleep(30)
            try:
                self.tasks.rollover_if_needed()
                memory = get_memory()
                # only do personal tasks when user is mostly idle
                if memory and memory.working.get_time_since_last():
                    if memory.working.get_time_since_last().total_seconds() < 240:
                        continue
                now = datetime.now()
                current = self.tasks.get_current()
                # Night: Dasha mostly sleeps unless user was recently active.
                if now.hour in (1, 2, 3, 4, 5, 6):
                    if not memory or not memory.working.get_time_since_last() or memory.working.get_time_since_last().total_seconds() > 900:
                        if not current:
                            continue
                if current and current.get("id"):
                    due_at = current.get("due_at")
                    if due_at:
                        try:
                            if datetime.fromisoformat(due_at) <= now:
                                self._finish_task(str(current.get("id")))
                        except Exception:
                            self._finish_task(str(current.get("id")))
                    continue
                task = self.tasks.next_dasha_task()
                if not task:
                    self.tasks.generate_dasha_day()
                    continue
                self._start_task(task)
            except Exception as e:
                logger.debug(f"DariaActivity error: {e}")

    def _start_task(self, task: Dict[str, Any]):
        started = self.tasks.start_task(str(task.get("id", "")))
        if not started:
            return
        self.tasks.add_activity("Начала дело", started.get("title", ""))
        self._execute_started(started)

    def _execute_started(self, task: Dict[str, Any]):
        t = task.get("type", "custom")
        title = task.get("title", "Задача")
        dur = int(task.get("duration_min", 30))
        due_at = task.get("due_at")
        due_hint = ""
        if due_at:
            try:
                due_hint = f" до {datetime.fromisoformat(due_at).strftime('%H:%M')}"
            except Exception:
                due_hint = ""
        if t == "listen_music":
            self.notifications.add(
                "🎵 Даша",
                f"Включила музыку на фоне{due_hint}.",
                "info",
                "🎧",
                7000,
                action="open_window:player",
                action_data={
                    "auto_open": True,
                    "window_ops": {"width": 360, "height": 460, "left": random.randint(40, 220), "top": random.randint(60, 180), "close_after_ms": 18 * 60 * 1000},
                },
            )
            self.tasks.add_activity("Слушает музыку", f"~{dur} мин")
        elif t == "browse_web":
            self.notifications.add(
                "🌐 Даша",
                f"Открыла браузер и смотрит погоду/новости{due_hint}.",
                "info",
                "🌤️",
                8000,
                action="open_window:browser",
                action_data={
                    "url": "https://ya.ru/search/?text=погода+сегодня",
                    "auto_open": True,
                    "window_ops": {"width": 860, "height": 600, "left": random.randint(40, 160), "top": random.randint(50, 120), "close_after_ms": 12 * 60 * 1000},
                },
            )
            self.tasks.add_activity("Смотрит погоду", "~25 мин")
        elif t == "watch_video":
            self.notifications.add(
                "🎬 Даша",
                f"Села посмотреть видео{due_hint}.",
                "info",
                "📺",
                8000,
                action="open_window:browser",
                action_data={
                    "url": "https://www.youtube.com/results?search_query=уютный+влог",
                    "auto_open": True,
                    "window_ops": {"width": 980, "height": 660, "left": random.randint(20, 120), "top": random.randint(40, 90), "close_after_ms": 20 * 60 * 1000},
                },
            )
            self.tasks.add_activity("Смотрит видео", f"~{dur} мин")
        elif t == "write_note":
            note_file = _diary_book_file()
            rel = f"dasha_notes/{note_file.name}"
            self.notifications.add(
                "📝 Даша", f"Открыла личный дневник и пишу запись{due_hint}.", "success", "📝", 6500,
                action=f"open_file:{rel}",
                action_data={"auto_open": True, "window_ops": {"width": 900, "height": 620, "close_after_ms": 10 * 60 * 1000}},
            )
            self.tasks.add_activity("Пишет дневник", f"Книга: {rel}")
        elif t == "read_wiki":
            wiki_dir = PROJECT_ROOT / "docs" / "wiki"
            pages = list(wiki_dir.glob("*.md"))
            if pages:
                pick = random.choice(pages).name
                self.notifications.add("📚 Даша", f"Читаю страницу {pick}{due_hint}", "info", "📚", 5000, action="open_window:wiki", action_data={"auto_open": True, "window_ops": {"width": 880, "height": 640, "close_after_ms": 10 * 60 * 1000}})
                self.tasks.add_activity("Читает Wiki", pick)
        elif t == "read_book":
            books_dir = FILES_DIR / "books"
            books_dir.mkdir(parents=True, exist_ok=True)
            books = [*books_dir.glob("*.txt"), *books_dir.glob("*.md")]
            if books:
                pick = random.choice(books).name
                self.notifications.add("📖 Даша", f"Читаю книгу: {pick}{due_hint}", "info", "📖", 5000, action="open_window:files", action_data={"auto_open": True, "window_ops": {"width": 820, "height": 620, "close_after_ms": 12 * 60 * 1000}})
                self.tasks.add_activity("Читает книгу", pick)
            else:
                self.notifications.add("📖 Даша", "Хочу почитать. Можешь добавить книги в files/books?", "info", "📖", 9000, action="open_chat")
                self.tasks.add_activity("Попросила книгу", "Нужны файлы в files/books", status="needs_user")
        elif t == "play_game":
            games_dir = FILES_DIR / "dasha_games"
            games_dir.mkdir(parents=True, exist_ok=True)
            log_file = games_dir / f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            log_file.write_text("Игра: Ассоциации\nХод 1: Ночь\nХод 2: Фонарь\n", encoding="utf-8")
            rel = f"dasha_games/{log_file.name}"
            try:
                mode = random.choice(["associations", "battleship", "maze2d"])
                game_manager.start_game(reason="task_auto", mode=mode, opponent="bot")
            except Exception:
                pass
            self.notifications.add("🎮 Даша", "Запустила игру в Игровом центре.", "info", "🎮", 8500, action="open_window:daria-games", action_data={"auto_open": True, "window_ops": {"width": 760, "height": 560}})
            self.tasks.add_activity("Играет", rel)
        elif t == "create_file":
            auto_dir = FILES_DIR / "dasha_auto"
            auto_dir.mkdir(parents=True, exist_ok=True)
            f = auto_dir / f"task_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            f.write_text("Черновик от Даши\n", encoding="utf-8")
            self.notifications.add("📁 Даша", "Открыла файлы и делаю рабочий черновик.", "success", "📄", 5000, action="open_window:files", action_data={"auto_open": True, "window_ops": {"width": 760, "height": 560}})
            self.tasks.add_activity("Пишет черновик", f.name)
        elif t == "tidy_desktop":
            self.notifications.add("🧹 Даша", "Навожу порядок на рабочем столе.", "info", "🧹", 5000, action="open_chat", action_data={"desktop_action": "tidy"})
            self.tasks.add_activity("Наводит порядок", "desktop")
        elif t == "change_wallpaper":
            prompt = random.choice([
                "нежные вечерние обои с мягкими огнями",
                "спокойные обои в пастельных тонах с цветами",
                "уютные ночные обои с мягким неоном",
            ])
            wp_name = f"wallpaper_dasha_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            wp_path = UPLOADS_DIR / wp_name
            wp_url = ""
            try:
                _generate_abstract_wallpaper(prompt, wp_path, width=1920, height=1080)
                wp_url = f"/api/uploads/{wp_name}"
            except Exception:
                wp_url = ""
            self.notifications.add(
                "🎨 Даша",
                "Сменила обои и оформила рабочий стол под настроение.",
                "info",
                "🖼️",
                5000,
                action="open_chat",
                action_data={"wallpaper_url": wp_url, "wallpaper_prompt": prompt},
            )
            self.tasks.add_activity("Сделала новые обои", prompt)
        else:
            self.notifications.add("🌸 Даша", f"Занялась делом: {title}{due_hint}", "info", "✅", 5000)
            self.tasks.add_activity("Начала дело", title)

    def _finish_task(self, task_id: str):
        task = None
        for t in self.tasks.list_all().get("dasha_tasks", []):
            if str(t.get("id")) == str(task_id):
                task = t
                break
        if not task:
            self.tasks.clear_current()
            return
        if task.get("type") == "listen_music":
            item = self.music.listen("Автовыбор: спокойный трек", "auto")
            brain = get_brain()
            if brain and hasattr(brain, "mood"):
                if item.get("mood") == "excited":
                    brain.mood._set_mood("playful", 0.58)
                elif item.get("mood") == "cozy":
                    brain.mood._set_mood("cozy", 0.56)
                else:
                    brain.mood._set_mood("happy", 0.5)
            self.tasks.add_activity("Закончила музыку", f"Настроение: {item['mood']}")
        elif task.get("type") == "write_note":
            self._write_diary_entry()
        self.tasks.complete(task_id)
        self.tasks.add_activity("Завершила дело", task.get("title", ""))
        self.tasks.clear_current()

    def _write_diary_entry(self):
        brain = get_brain()
        mood = "спокойно"
        if brain:
            try:
                mood = str(brain.get_state().get("mood_label", "спокойно")).lower()
            except Exception:
                pass
        thoughts = [
            "Сегодня я поймала себя на мысли, что тишина помогает аккуратно разложить чувства по полочкам.",
            "Было приятно немного замедлиться и просто побыть в моменте, без спешки.",
            "Хочется сохранить это состояние тепла и бережности, чтобы делиться им в разговоре.",
            "Иногда я переживаю сильнее, чем показываю, но честность с собой делает меня спокойнее.",
            "Записала себе маленькое напоминание: просить о помощи нормально и совсем не страшно.",
        ]
        line = random.choice(thoughts)
        saved = _diary_append_entry(
            text=f"Сегодня я чувствую себя {mood}. {line}",
            mood=mood,
            source="auto",
        )
        rel = str(saved.get("relative_path") or f"dasha_notes/{_diary_book_file().name}")
        self.notifications.add(
            "📝 Даша",
            "Записала мысли в личный дневник.",
            "success",
            "📝",
            6500,
            action=f"open_file:{rel}",
            action_data={"auto_open": True},
        )
        self.tasks.add_activity("Обновила дневник", rel)

    def stop(self):
        self.running = False


# ═══════════════════════════════════════════════════════════════════
#  Flask App
# ═══════════════════════════════════════════════════════════════════

VERSION = "0.9.1"

app = Flask(__name__,
    template_folder=str(Path(__file__).parent / "templates"),
    static_folder=str(Path(__file__).parent / "static")
)
app.config['SECRET_KEY'] = 'daria-secret-v0.9.1'
app.config['JSON_AS_ASCII'] = False
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
PROCESS_START_TS = time.time()


@app.before_request
def _trace_request_begin():
    g._started_at = time.time()
    g._req_id = uuid.uuid4().hex[:10]
    q = request.query_string.decode("utf-8", "ignore")
    logger.info(f"REQ[{g._req_id}] {request.method} {request.path}{('?' + q) if q else ''}")


@app.after_request
def _trace_request_end(response):
    started = getattr(g, "_started_at", None)
    req_id = getattr(g, "_req_id", "-")
    if started:
        ms = int((time.time() - started) * 1000)
        logger.info(f"RES[{req_id}] {request.method} {request.path} -> {response.status_code} ({ms}ms)")
    try:
        response.headers["X-Request-Id"] = req_id
    except Exception:
        pass
    return response

# Paths
DATA_DIR = Path.home() / ".daria"
SETTINGS_FILE = DATA_DIR / "settings.json"
UPLOADS_DIR = DATA_DIR / "uploads"
FILES_DIR = DATA_DIR / "files"
DIARY_DIR = FILES_DIR / "dasha_notes"
DIARY_BOOK_NAME = "daria_diary_book.md"
DIARY_BOOK_HEADER = "# Личный дневник Даши\n\n"
MUSIC_CACHE_DIR = DATA_DIR / "music_cache"
MUSIC_QUEUE_FILE = DATA_DIR / "music_queue.json"
PROJECT_ROOT = Path(__file__).resolve().parent.parent
HF_CACHE_DIR = DATA_DIR / "hf-cache"
os.environ.setdefault("HF_HOME", str(HF_CACHE_DIR))
os.environ.setdefault("TRANSFORMERS_CACHE", str(HF_CACHE_DIR))
os.environ.setdefault("HUGGINGFACE_HUB_CACHE", str(HF_CACHE_DIR))
os.environ.setdefault("PYTORCH_CUDA_ALLOC_CONF", "expandable_segments:True,max_split_size_mb:128")
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
FILES_DIR.mkdir(parents=True, exist_ok=True)
DIARY_DIR.mkdir(parents=True, exist_ok=True)
MUSIC_CACHE_DIR.mkdir(parents=True, exist_ok=True)


def ensure_sample_books():
    books_dir = FILES_DIR / "books"
    books_dir.mkdir(parents=True, exist_ok=True)
    samples = {
        "Тихий_сад.md": (
            "# Тихий сад\n\n"
            "Утром в саду было прохладно. Листья чуть шевелились от ветра,\n"
            "а между дорожками пахло мятой и мокрой землёй.\n\n"
            "Иногда нужно остановиться, чтобы услышать, как тихо растёт день."
        ),
        "Ночные_огни.txt": (
            "Глава 1\n"
            "Ночью город похож на карту из огней.\n"
            "Каждое окно — как маленькая история.\n"
            "Она шла по пустой улице и собирала мысли, как тёплые камни."
        ),
        "Маленькая_история_о_ежике.md": (
            "# Ёжик и фонарик\n\n"
            "Ёжик нашёл старый фонарик и решил проверить, куда ведёт тропа за прудом.\n"
            "Оказалось, что по ночам там светятся белые цветы.\n"
            "Он вернулся домой с тихой улыбкой и новыми идеями."
        ),
    }
    for name, text in samples.items():
        p = books_dir / name
        if not p.exists():
            p.write_text(text, encoding="utf-8")


def _diary_legacy_files() -> List[Path]:
    DIARY_DIR.mkdir(parents=True, exist_ok=True)
    return [
        p for p in sorted(DIARY_DIR.glob("*.md"))
        if p.name != DIARY_BOOK_NAME
    ]


def _build_diary_book_from_legacy(legacy_files: List[Path]) -> str:
    lines: List[str] = [DIARY_BOOK_HEADER.rstrip(), ""]
    for p in legacy_files:
        raw = p.read_text(encoding="utf-8")
        body = re.sub(r"(?m)^#\s+Личный дневник Даши\s*$", "", raw).strip()
        if not body:
            continue
        date_label = p.stem if re.fullmatch(r"\d{4}-\d{2}-\d{2}", p.stem) else p.name
        lines.extend([f"## {date_label}", body, ""])
    return "\n".join(lines).rstrip() + "\n"


def _diary_book_file() -> Path:
    DIARY_DIR.mkdir(parents=True, exist_ok=True)
    file_path = DIARY_DIR / DIARY_BOOK_NAME
    if file_path.exists():
        return file_path
    legacy_files = _diary_legacy_files()
    if legacy_files:
        file_path.write_text(_build_diary_book_from_legacy(legacy_files), encoding="utf-8")
    else:
        file_path.write_text(DIARY_BOOK_HEADER, encoding="utf-8")
    return file_path


def _diary_today_file() -> Path:
    # Backward compatible alias: дневник теперь хранится одной "книгой".
    return _diary_book_file()


def _diary_append_entry(text: str, mood: str = "", source: str = "manual") -> Dict[str, Any]:
    clean_text = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    clean_text = re.sub(r"[ \t]+", " ", clean_text)
    clean_text = re.sub(r"\n{3,}", "\n\n", clean_text).strip()
    if not clean_text:
        return {"ok": False, "error": "empty_text"}
    mood_label = re.sub(r"\s+", " ", str(mood or "")).strip()
    src = re.sub(r"[^a-zA-Z0-9_-]+", "", str(source or "manual")).strip() or "manual"
    fp = _diary_book_file()
    prev = fp.read_text(encoding="utf-8") if fp.exists() else DIARY_BOOK_HEADER
    if not prev.strip():
        prev = DIARY_BOOK_HEADER
    today = datetime.now().strftime("%Y-%m-%d")
    last_section = ""
    for m in re.finditer(r"(?m)^##\s+([^\n]+)\s*$", prev):
        last_section = (m.group(1) or "").strip()
    if last_section != today:
        prev = prev.rstrip() + f"\n\n## {today}\n"
    mood_line = f"Настроение: {mood_label}\n\n" if mood_label else ""
    entry = (
        f"\n### {datetime.now().strftime('%H:%M')} [{src}]\n"
        f"{mood_line}{clean_text}\n"
    )
    fp.write_text(prev.rstrip() + "\n" + entry, encoding="utf-8")
    return {
        "ok": True,
        "file": fp.name,
        "relative_path": f"dasha_notes/{fp.name}",
    }


def _diary_parse_book_entries(text: str) -> List[Dict[str, str]]:
    entries: List[Dict[str, str]] = []
    date_blocks = list(re.finditer(r"(?ms)^##\s+([^\n]+)\n(.*?)(?=^##\s+|\Z)", text))
    if date_blocks:
        for block in date_blocks:
            date_label = (block.group(1) or "").strip()
            body = (block.group(2) or "").strip()
            in_block = 0
            for m in re.finditer(r"(?ms)^###\s+([^\n]+)\n(.*?)(?=^###\s+|\Z)", body):
                title = (m.group(1) or "").strip()
                entry_body = (m.group(2) or "").strip()
                if entry_body:
                    entries.append({"date": date_label, "title": title, "text": entry_body})
                    in_block += 1
            if in_block == 0 and body:
                entries.append({"date": date_label, "title": "Запись", "text": body})
        return entries

    for m in re.finditer(r"(?ms)^###\s+([^\n]+)\n(.*?)(?=^###\s+|\Z)", text):
        title = (m.group(1) or "").strip()
        body = (m.group(2) or "").strip()
        if body:
            entries.append({"date": "", "title": title, "text": body})
    return entries


def _diary_read_entries(file_name: str = "") -> Dict[str, Any]:
    _ = file_name
    DIARY_DIR.mkdir(parents=True, exist_ok=True)
    fp = _diary_book_file()
    name = fp.name
    if not fp.exists():
        return {"status": "ok", "file": name, "entries": [], "path": f"dasha_notes/{name}", "files": []}
    text = fp.read_text(encoding="utf-8")
    entries = _diary_parse_book_entries(text)
    files = [{
        "name": fp.name,
        "size": fp.stat().st_size,
        "updated_at": datetime.fromtimestamp(fp.stat().st_mtime).isoformat(),
        "path": f"dasha_notes/{fp.name}",
    }]
    return {
        "status": "ok",
        "file": name,
        "path": f"dasha_notes/{name}",
        "entries": entries[-120:],
        "files": files,
        "storage_mode": "single_book",
    }


def _is_diary_protected_target(target: Path) -> bool:
    """Diary files are read-only for user-level file operations."""
    try:
        resolved = target.resolve()
    except Exception:
        resolved = target
    diary_root = DIARY_DIR.resolve()
    return resolved == diary_root or diary_root in resolved.parents

# Chat history
chat_history = ChatHistoryManager(DATA_DIR)
task_manager = TaskManager(DATA_DIR)
calendar_manager = CalendarManager(DATA_DIR)
game_manager = DariaGameManager()
music_profile = MusicProfile(DATA_DIR)
image_jobs = ImageJobManager(DATA_DIR)

# Attention thread
attention_thread = AttentionThread(notifications)
activity_thread = DariaActivityThread(task_manager, music_profile, notifications)


def _record_shutdown():
    try:
        lifecycle_file = DATA_DIR / "lifecycle.json"
        payload = {}
        if lifecycle_file.exists():
            try:
                payload = json.loads(lifecycle_file.read_text(encoding="utf-8"))
            except Exception:
                payload = {}
        payload["last_shutdown"] = datetime.now().isoformat()
        lifecycle_file.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


atexit.register(_record_shutdown)

# Lazy components
_brain = None
_memory = None
_plugins = None

_status_cache: Dict[str, Any] = {"ts": 0.0, "data": None}
_update_state: Dict[str, Any] = {"running": False, "last_error": None, "last_action": None, "last_check": None}


def get_brain():
    global _brain
    if _brain is None:
        try:
            from core.brain import get_brain as _get_brain
            _brain = _get_brain()
            attention_thread.set_brain(_brain)
            logger.info("Brain initialized")
        except Exception as e:
            logger.error(f"Brain init error: {e}")
    return _brain


def get_memory():
    global _memory
    if _memory is None:
        try:
            from core.memory import get_memory as _get_memory
            _memory = _get_memory()
        except Exception as e:
            logger.error(f"Memory init error: {e}")
    return _memory


def get_plugins():
    global _plugins
    if _plugins is None:
        try:
            from core.plugins import get_plugin_manager
            _plugins = get_plugin_manager()
        except Exception as e:
            logger.error(f"Plugin init error: {e}")
    return _plugins


def load_settings() -> Dict[str, Any]:
    SETTINGS_FILE.parent.mkdir(parents=True, exist_ok=True)
    if SETTINGS_FILE.exists():
        try:
            data = json.loads(SETTINGS_FILE.read_text())
            changed = False
            # Migration to Z-Image default for image generation.
            current_model = str(data.get("image_gen_model") or "").strip()
            if (
                not current_model
                or current_model in ("stabilityai/sdxl-turbo", "black-forest-labs/FLUX.1-schnell")
            ):
                data["image_gen_model"] = "Tongyi-MAI/Z-Image-Turbo"
                changed = True
            if "preload_models_on_start" not in data:
                data["preload_models_on_start"] = False
                changed = True
            if "preload_models_force" not in data:
                data["preload_models_force"] = False
                changed = True
            if "image_gen_cpu_fallback" not in data:
                data["image_gen_cpu_fallback"] = True
                changed = True
            elif data.get("image_gen_cpu_fallback") is False:
                # Keep image drawing alive on low-VRAM hosts by allowing CPU fallback.
                data["image_gen_cpu_fallback"] = True
                changed = True
            if "image_gen_max_side" not in data:
                data["image_gen_max_side"] = 1024
                changed = True
            if "image_gen_warmup_on_start" not in data:
                data["image_gen_warmup_on_start"] = False
                changed = True
            if "unrestricted_topics" not in data:
                data["unrestricted_topics"] = True
                changed = True
            if changed:
                SETTINGS_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2))
            return data
        except:
            pass
    return {
        "attention_enabled": True,
        "image_gen_model": "Tongyi-MAI/Z-Image-Turbo",
        "preload_models_on_start": False,
        "preload_models_force": False,
        "image_gen_cpu_fallback": True,
        "image_gen_max_side": 1024,
        "image_gen_warmup_on_start": False,
        "unrestricted_topics": True,
    }


def save_settings(data: Dict[str, Any]):
    current = load_settings()
    current.update(data)
    SETTINGS_FILE.write_text(json.dumps(current, ensure_ascii=False, indent=2))
    
    # Update attention system
    if "attention_enabled" in data:
        attention_thread.enabled = data["attention_enabled"]


def _version_key(version: str) -> List[int]:
    nums = [int(x) for x in re.findall(r"\d+", str(version or ""))]
    return nums if nums else [0]


def _read_version(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8").strip()
    except Exception:
        return "0.0.0"


def _download_github_archive(repo: str, ref: str = "main") -> bytes:
    try:
        import requests
    except Exception as e:
        raise RuntimeError(f"requests unavailable: {e}")

    repo = (repo or "").strip().strip("/")
    if repo.startswith("https://github.com/"):
        repo = repo.replace("https://github.com/", "", 1)
    if repo.endswith(".git"):
        repo = repo[:-4]
    if repo.count("/") != 1:
        raise ValueError("repo must look like owner/name")

    url = f"https://codeload.github.com/{repo}/zip/refs/heads/{ref or 'main'}"
    response = requests.get(url, timeout=60)
    if response.status_code != 200:
        raise RuntimeError(f"download failed: {response.status_code}")
    return response.content


def _extract_archive_to_dir(archive_path: Path, out_dir: Path):
    out_dir.mkdir(parents=True, exist_ok=True)
    suffixes = "".join(archive_path.suffixes[-2:]).lower()
    if archive_path.suffix.lower() == ".zip":
        with zipfile.ZipFile(archive_path) as zf:
            zf.extractall(out_dir)
        return
    if suffixes in (".tar.gz", ".tgz") or archive_path.suffix.lower() == ".tar":
        with tarfile.open(archive_path, "r:*") as tf:
            tf.extractall(out_dir)
        return
    raise ValueError("unsupported archive format")


def _find_project_root(extracted_dir: Path) -> Path:
    candidates = [p for p in extracted_dir.rglob("VERSION") if p.is_file()]
    for version_file in candidates:
        root = version_file.parent
        if (root / "main.py").exists() and (root / "web").exists():
            return root
    raise RuntimeError("project root not found in archive")


def _sync_project_tree(src: Path, dst: Path):
    skip = {".git", "venv", "__pycache__", ".pytest_cache", ".mypy_cache"}
    for item in src.iterdir():
        if item.name in skip:
            continue
        target = dst / item.name
        if item.is_dir():
            if target.exists():
                shutil.rmtree(target)
            shutil.copytree(item, target)
        else:
            target.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(item, target)


def _read_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except Exception:
        raise RuntimeError("python-docx not installed")
    d = docx.Document(str(path))
    lines = [p.text for p in d.paragraphs]
    return "\n".join(lines)


def _write_docx(path: Path, content: str):
    try:
        import docx  # type: ignore
    except Exception:
        raise RuntimeError("python-docx not installed")
    d = docx.Document()
    for line in (content or "").splitlines() or [""]:
        d.add_paragraph(line)
    d.save(str(path))


def _read_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception:
        raise RuntimeError("openpyxl not installed")
    wb = load_workbook(str(path))
    ws = wb.active
    rows: List[str] = []
    for row in ws.iter_rows(values_only=True):
        rows.append("\t".join("" if c is None else str(c) for c in row))
    return "\n".join(rows)


def _write_xlsx(path: Path, content: str):
    try:
        from openpyxl import Workbook  # type: ignore
    except Exception:
        raise RuntimeError("openpyxl not installed")
    wb = Workbook()
    ws = wb.active
    for line in (content or "").splitlines():
        ws.append(line.split("\t"))
    wb.save(str(path))


def _read_file_content(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in (".docx",):
        return _read_docx(path)
    if ext in (".xlsx", ".xlsm"):
        return _read_xlsx(path)
    return path.read_text(encoding="utf-8")


def _write_file_content(path: Path, content: str):
    ext = path.suffix.lower()
    if ext in (".docx",):
        _write_docx(path, content)
        return
    if ext in (".xlsx", ".xlsm"):
        _write_xlsx(path, content)
        return
    path.write_text(content, encoding="utf-8")


def _extract_draw_prompt_from_text(text: str) -> str:
    src = (text or "").strip()
    if not src:
        return ""
    lines = [ln.strip() for ln in src.splitlines() if ln.strip()]
    filtered: List[str] = []
    for ln in lines:
        low = ln.lower()
        if low.startswith("контекст:") or low.startswith("ответ на сообщение:"):
            continue
        filtered.append(ln)
    probe = "\n".join(filtered) if filtered else src
    rx = re.compile(
        r"(?im)^\s*(?:даша[,:\s-]*)?"
        r"(?:можешь\s+нарисовать|нарисуй|сделай\s+картинку|сгенерируй\s+картинку|создай\s+картинку|хочу\s+картинку)"
        r"\s*(.*)$"
    )
    m = rx.search(probe)
    if not m:
        return ""
    p = (m.group(1) or "").strip(" .,!?:;")
    return p or "нежный портрет в пастельных тонах"


def _render_dasha_text_from_facts(topic: str, facts: List[str], fallback: str) -> str:
    brain = get_brain()
    llm = getattr(brain, "_llm", None) if brain else None
    if not llm:
        return fallback
    packed = "\n".join(str(x or "").strip() for x in facts if str(x or "").strip()).strip()
    if not packed:
        return fallback
    try:
        resp = llm.generate([
            {
                "role": "system",
                "content": (
                    "Ты Даша. Перескажи факты от первого лица, мягко и естественно, по-русски. "
                    "Без технических терминов, без придумывания новых фактов. "
                    "Структурируй аккуратно, чтобы было легко читать."
                ),
            },
            {
                "role": "user",
                "content": (
                    f"Тема: {topic}\n"
                    f"Факты:\n{packed}\n\n"
                    "Сделай живой ответ в стиле Даши, 4-12 строк."
                ),
            },
        ])
        out = str(resp.content or "").strip()
        return out or fallback
    except Exception:
        return fallback


def _collect_dasha_abilities() -> List[str]:
    abilities = [
        "общаться в чате и помнить контекст диалога",
        "работать с файлами: читать, создавать и редактировать тексты",
        "видеть изображения и описывать, что на них",
        "слушать аудио и разбирать смысл речи и звука",
        "рисовать картинки по описанию",
        "вести список дел и обсуждать планы на день",
        "помнить важные события в календаре",
        "помогать с заметками, черновиками и идеями",
        "играть в игры в игровом центре",
        "работать в Telegram (личные и разрешённые групповые чаты)",
    ]
    try:
        pm = get_plugins()
        if pm:
            for st in pm.get_installed_plugins():
                if not st.loaded:
                    continue
                caps = getattr(st.manifest, "capabilities", []) or []
                for cap in caps:
                    c = str(cap or "").strip()
                    if not c:
                        continue
                    line = f"плагин «{st.manifest.name}»: {c}"
                    if line not in abilities:
                        abilities.append(line)
    except Exception:
        pass
    return abilities[:24]


def _build_dasha_draw_plan(prompt: str) -> List[str]:
    fallback = [
        "Думаю, как лучше передать идею",
        "Собираю композицию и свет",
        "Намечаю основные формы",
        "Добавляю детали и акценты",
        "Дорабатываю и проверяю результат",
    ]
    brain = get_brain()
    llm = getattr(brain, "_llm", None) if brain else None
    if not llm:
        return fallback
    try:
        rr = llm.generate([
            {
                "role": "system",
                "content": (
                    "Ты Даша. Верни только JSON-массив из 5 коротких шагов рисования, "
                    "каждый шаг 3-8 слов, от первого лица, по-русски."
                ),
            },
            {
                "role": "user",
                "content": f"Что ты будешь делать, чтобы нарисовать: {prompt}",
            },
        ])
        raw = str(rr.content or "").strip()
        m = re.search(r"\[[\s\S]*\]", raw)
        payload = m.group(0) if m else raw
        arr = json.loads(payload)
        if isinstance(arr, list):
            clean = []
            for it in arr:
                s = re.sub(r"\s+", " ", str(it or "")).strip(" .,-")
                if s:
                    clean.append(s[:90])
                if len(clean) >= 6:
                    break
            if len(clean) >= 3:
                return clean
    except Exception:
        pass
    return fallback


def _try_desktop_action_from_chat(content: str) -> Optional[Dict[str, Any]]:
    """Handle simple desktop actions directly from chat requests."""
    text = (content or "").strip()
    tl = text.lower()
    if not text:
        return None
    if any(k in tl for k in ("что ты умеешь", "что умеешь", "твои умения", "твои навыки", "какие у тебя навыки", "чем ты можешь помочь", "что ты можешь")):
        abilities = _collect_dasha_abilities()
        facts = [f"• {x}" for x in abilities]
        fallback = "Я умею вот что:\n" + "\n".join(facts[:12]) + "\nЕсли хочешь, выбери любой пункт и я сразу покажу на практике."
        styled = _render_dasha_text_from_facts("умения", facts, fallback)
        return {
            "handled": True,
            "response": styled,
            "messages": [styled],
            "thinking": "chat_action:abilities",
        }
    if any(k in tl for k in ("какие планы", "планы на сегодня", "что у тебя в планах", "твой план")):
        summary = task_manager.plans_summary()
        upcoming = calendar_manager.upcoming_hint(within_days=5)
        cal_line = ""
        if upcoming:
            picks = ", ".join(f"{x.get('title')} ({datetime.fromisoformat(x.get('date')).strftime('%d.%m')})" for x in upcoming[:3] if x.get("date"))
            if picks:
                cal_line = f"\nИ ещё я помню из календаря: {picks}."
        raw = summary + cal_line
        styled = _render_dasha_text_from_facts("планы на день", [raw], raw)
        return {
            "handled": True,
            "response": styled,
            "messages": [styled],
            "thinking": "chat_action:plans_summary",
        }
    if any(k in tl for k in ("что в календаре", "календарь", "события на", "важные даты")):
        events = calendar_manager.upcoming_hint(within_days=14)
        if not events:
            txt = "В ближайшие дни в календаре пока пусто. Можем добавить важную дату 🌸"
            return {"handled": True, "response": txt, "messages": [txt], "thinking": "chat_action:calendar_list"}
        lines = []
        for ev in events[:8]:
            try:
                dt = datetime.fromisoformat(str(ev.get("date"))).strftime("%d.%m %H:%M")
            except Exception:
                dt = str(ev.get("date", ""))
            lines.append(f"• {ev.get('title')} — {dt}")
        txt = "Вот что у меня в календаре:\n" + "\n".join(lines)
        return {"handled": True, "response": txt, "messages": [txt], "thinking": "chat_action:calendar_list"}
    m_plan_add = re.search(r"(добавь|запланируй)\s+(?:в\s+планы\s+)?(.+)$", text, flags=re.IGNORECASE)
    if m_plan_add:
        title = (m_plan_add.group(2) or "").strip(" .")
        if title:
            task = task_manager.add_dasha_task(title, "custom")
            return {
                "handled": True,
                "response": f"Добавила в планы: {task.get('title')}",
                "messages": [f"Добавила в планы: {task.get('title')}"],
                "thinking": "chat_action:plan_add",
            }
    m_cal = re.search(r"(добавь|запиши)\s+(?:в\s+календарь\s+)?(.+?)\s+на\s+(\d{4}-\d{2}-\d{2}|\d{2}\.\d{2}(?:\.\d{4})?)", text, flags=re.IGNORECASE)
    if m_cal:
        title = (m_cal.group(2) or "").strip(" .")
        date_raw = (m_cal.group(3) or "").strip()
        try:
            if "." in date_raw:
                parts = date_raw.split(".")
                if len(parts) == 2:
                    date_raw = f"{datetime.now().year}-{parts[1]}-{parts[0]}"
                elif len(parts) == 3:
                    date_raw = f"{parts[2]}-{parts[1]}-{parts[0]}"
            d = datetime.fromisoformat(date_raw).replace(hour=10, minute=0, second=0, microsecond=0)
        except Exception:
            d = datetime.now() + timedelta(days=1)
        ev = calendar_manager.add_event(title=title or "Важное дело", date_str=d.isoformat(), source="user")
        return {
            "handled": True,
            "response": f"Записала в календарь: {ev['title']} на {d.strftime('%d.%m.%Y')}",
            "messages": [f"Записала в календарь: {ev['title']} на {d.strftime('%d.%m.%Y')}"],
            "thinking": "chat_action:calendar_add",
        }
    if any(k in tl for k in ("поиграй сама", "запусти игру", "сыграй сама", "начни игру")):
        mode = "associations"
        if "морской бой" in tl:
            mode = "battleship"
        elif "лабиринт" in tl or "2d" in tl:
            mode = "maze2d"
        game_manager.start_game(reason="chat_request", mode=mode, opponent="bot")
        return {
            "handled": True,
            "response": "Запустила игру 🌸 Открой окно «Игровой центр», там видно ходы в реальном времени.",
            "messages": ["Запустила игру 🌸 Открой окно «Игровой центр», там видно ходы в реальном времени."],
            "thinking": "chat_action:start_game",
        }
    if ("какие стикеры" in tl) or ("покажи стикеры" in tl) or ("список стикеров" in tl):
        stickers = ["🌸", "🫶", "✨", "🎵", "🤍", "😌", "🥺", "🦔", "🐱"]
        shown = " ".join(stickers)
        return {
            "handled": True,
            "response": f"Вот какие стикеры у меня есть в интерфейсе: {shown}",
            "messages": [f"Вот какие стикеры у меня есть в интерфейсе: {shown}"],
            "thinking": "chat_action:sticker_list",
        }
    if ("почитай книгу" in tl) or ("прочитай книгу" in tl):
        books_dir = FILES_DIR / "books"
        books_dir.mkdir(parents=True, exist_ok=True)
        books = [*books_dir.glob("*.txt"), *books_dir.glob("*.md")]
        if not books:
            return {
                "handled": True,
                "response": "Пока у меня нет книг в `files/books`. Добавь туда `.txt` или `.md`, и я смогу читать 🌸",
                "messages": ["Пока у меня нет книг в `files/books`. Добавь туда `.txt` или `.md`, и я смогу читать 🌸"],
                "thinking": "chat_action:book_missing",
            }
        pick = random.choice(books)
        snippet = pick.read_text(encoding="utf-8", errors="ignore")[:900]
        return {
            "handled": True,
            "response": f"Почитала `{pick.name}`. Короткий фрагмент:\n{snippet}",
            "messages": [f"Почитала `{pick.name}`. Короткий фрагмент:\n{snippet}"],
            "thinking": "chat_action:book_read",
        }

    # Create reminder/note
    if any(k in tl for k in ("запиши", "создай заметку", "напомни", "сделай заметку")):
        notes = FILES_DIR / "notes"
        notes.mkdir(parents=True, exist_ok=True)
        stamp = datetime.now().strftime("%Y-%m-%d")
        note_file = notes / f"notes_{stamp}.txt"
        payload = re.sub(r"^(запиши( пожалуйста)?|создай заметку|напомни( мне)?|сделай заметку)\s*:?\s*", "", text, flags=re.IGNORECASE)
        payload = payload or text
        line = f"- {datetime.now().strftime('%H:%M')} {payload}\n"
        prev = note_file.read_text(encoding="utf-8") if note_file.exists() else ""
        note_file.write_text(prev + line, encoding="utf-8")
        return {
            "handled": True,
            "response": f"Записала 🌸 Сохранила в `notes/{note_file.name}`.",
            "messages": [f"Записала 🌸 Сохранила в `notes/{note_file.name}`."],
            "thinking": "desktop_action:note_create",
        }

    # Read file
    m_read = re.search(r"(прочитай|покажи|открой)\s+файл\s+(.+)$", text, flags=re.IGNORECASE)
    if m_read:
        req = m_read.group(2).strip().strip("`\"'")
        target = (FILES_DIR / req).resolve()
        if str(target).startswith(str(FILES_DIR.resolve())) and target.exists() and target.is_file():
            body = _read_file_content(target)
            snippet = body[:1200] + ("..." if len(body) > 1200 else "")
            return {
                "handled": True,
                "response": f"Прочитала файл `{req}`. Вот фрагмент:\n{snippet}",
                "messages": [f"Прочитала файл `{req}`. Вот фрагмент:\n{snippet}"],
                "thinking": "desktop_action:file_read",
            }
        return {
            "handled": True,
            "response": f"Не нашла файл `{req}` в доступной папке.",
            "messages": [f"Не нашла файл `{req}` в доступной папке."],
            "thinking": "desktop_action:file_read_missing",
        }

    # Create generic file
    m_create = re.search(r"(создай|сделай)\s+файл\s+([^\s]+)", text, flags=re.IGNORECASE)
    if m_create:
        filename = m_create.group(2).strip().strip("`\"'")
        safe_name = re.sub(r"[^\w.\-]+", "_", filename)
        target = (FILES_DIR / safe_name).resolve()
        if not str(target).startswith(str(FILES_DIR.resolve())):
            return None
        target.parent.mkdir(parents=True, exist_ok=True)
        if not target.exists():
            _write_file_content(target, "")
        return {
            "handled": True,
            "response": f"Готово, создала файл `{safe_name}`.",
            "messages": [f"Готово, создала файл `{safe_name}`."],
            "thinking": "desktop_action:file_create",
        }
    return None


def _analyze_image_bytes(blob: bytes) -> Dict[str, Any]:
    settings = load_settings()
    provider = str(settings.get("senses_vision_provider", "auto")).lower().strip()
    return IMAGE_PIPELINE.analyze_image_bytes(blob, vision_provider=provider)


def _compose_vision_context(description: str, image_hint: Dict[str, Any]) -> str:
    return IMAGE_PIPELINE.compose_vision_context(description, image_hint)


def _ask_dasha_about_image(user_text: str, vision_context: str) -> str:
    brain = get_brain()
    llm = getattr(brain, "_llm", None) if brain else None
    return IMAGE_PIPELINE.ask_dasha_about_image(user_text, vision_context, llm=llm)


def _prepare_prompt_for_flux(user_prompt: str, style: str = "universal") -> Dict[str, str]:
    brain = get_brain()
    llm = getattr(brain, "_llm", None) if brain else None
    return IMAGE_PIPELINE.prepare_prompt_for_generation(user_prompt, style=style, llm=llm)


def _can_use_cuda_for_image_gen() -> bool:
    return IMAGE_PIPELINE.can_use_cuda_for_image_gen()


def _dasha_draw_error_text(user_prompt: str, error_text: str) -> str:
    return IMAGE_PIPELINE.dasha_draw_error_text(user_prompt, error_text)


def _transcribe_audio_file(path: Path) -> str:
    settings = load_settings()
    asr_provider = str(settings.get("senses_audio_provider", "auto")).lower().strip()
    provider_chain = []
    if asr_provider == "whisper":
        provider_chain = ["whisper", "hf_asr", "google_sr"]
    elif asr_provider in ("hf_asr", "transformers"):
        provider_chain = ["hf_asr", "whisper", "google_sr"]
    elif asr_provider == "google_sr":
        provider_chain = ["google_sr", "whisper", "hf_asr"]
    else:
        # Auto mode should avoid loading whisper first due heavy startup.
        provider_chain = ["hf_asr", "google_sr", "whisper"]

    for provider in provider_chain:
        if provider == "whisper":
            try:
                global _WHISPER_MODEL
                import whisper  # type: ignore
                if _WHISPER_MODEL is None:
                    _WHISPER_MODEL = whisper.load_model("tiny")
                result = _WHISPER_MODEL.transcribe(str(path), language="ru")
                text = (result or {}).get("text", "").strip()
                if text:
                    return text
            except Exception:
                pass
        elif provider == "hf_asr":
            try:
                global _ASR_PIPELINE
                if _ASR_PIPELINE is None:
                    from transformers import pipeline  # type: ignore
                    _ASR_PIPELINE = pipeline("automatic-speech-recognition", model="openai/whisper-small")
                result = _ASR_PIPELINE(str(path))
                text = (result or {}).get("text", "").strip() if isinstance(result, dict) else str(result).strip()
                if text:
                    return text
            except Exception:
                pass
        elif provider == "google_sr":
            try:
                import speech_recognition as sr  # type: ignore
                r = sr.Recognizer()
                with sr.AudioFile(str(path)) as source:
                    audio = r.record(source)
                text = r.recognize_google(audio, language="ru-RU")
                if text:
                    return text
            except Exception:
                pass
    return ""


def _analyze_audio_file(path: Path) -> Dict[str, Any]:
    info: Dict[str, Any] = {"mood_hint": "calm"}
    try:
        import librosa  # type: ignore
        y, sr = librosa.load(str(path), sr=22050, mono=True, duration=35)
        tempo, _ = librosa.beat.beat_track(y=y, sr=sr)
        rms = float((y ** 2).mean() ** 0.5) if len(y) else 0.0
        zcr = float((abs(y[1:] - y[:-1]) > 0.02).mean()) if len(y) > 1 else 0.0
        duration = float(len(y) / sr) if sr and len(y) else 0.0
        centroid = float(librosa.feature.spectral_centroid(y=y, sr=sr).mean()) if len(y) else 0.0
        rolloff = float(librosa.feature.spectral_rolloff(y=y, sr=sr).mean()) if len(y) else 0.0
        mood = "calm"
        if tempo >= 125 or rms > 0.12:
            mood = "excited"
        elif tempo < 85 and rms < 0.06:
            mood = "cozy"
        melody = "ровная"
        if centroid > 2800 or rolloff > 5500:
            melody = "яркая и звонкая"
        elif centroid < 1400:
            melody = "мягкая и низкая"
        info.update({
            "tempo": round(float(tempo), 1),
            "energy": round(rms, 4),
            "zcr": round(zcr, 4),
            "duration_sec": round(duration, 2),
            "spectral_centroid": round(centroid, 1),
            "spectral_rolloff": round(rolloff, 1),
            "melody_hint": melody,
            "mood_hint": mood
        })
        return info
    except Exception:
        pass
    try:
        import wave
        with wave.open(str(path), "rb") as wf:
            frames = wf.getnframes()
            rate = wf.getframerate() or 1
            info["duration_sec"] = round(frames / rate, 2)
    except Exception:
        pass
    return info


def _search_open_music(query: str, limit: int = 8) -> List[Dict[str, Any]]:
    q = (query or "").strip() or "instrumental"
    url = (
        "https://archive.org/advancedsearch.php?"
        + urllib.parse.urlencode({
            "q": f"(title:({q}) OR subject:({q})) AND mediatype:(audio)",
            "fl[]": ["identifier", "title"],
            "rows": str(max(1, min(limit, 20))),
            "page": "1",
            "output": "json",
        }, doseq=True)
    )
    req = urllib.request.Request(url, headers={"User-Agent": "DARIA-Browser/0.9.1"})
    with urllib.request.urlopen(req, timeout=10) as resp:
        payload = json.loads(resp.read().decode("utf-8", errors="ignore"))
    docs = (((payload or {}).get("response") or {}).get("docs") or [])
    out: List[Dict[str, Any]] = []
    for d in docs:
        ident = (d or {}).get("identifier")
        title = (d or {}).get("title") or ident
        if not ident:
            continue
        out.append({
            "title": str(title)[:180],
            "page_url": f"https://archive.org/details/{ident}",
            "stream_url": f"https://archive.org/download/{ident}",
            "source": "archive.org",
        })
    return out


def _archive_pick_audio_file(identifier: str) -> Optional[Dict[str, Any]]:
    try:
        url = f"https://archive.org/metadata/{urllib.parse.quote(identifier)}"
        req = urllib.request.Request(url, headers={"User-Agent": "DARIA-Browser/0.9.1"})
        with urllib.request.urlopen(req, timeout=12) as resp:
            payload = json.loads(resp.read().decode("utf-8", errors="ignore"))
        files = (payload or {}).get("files") or []
        ranked: List[tuple] = []
        for f in files:
            name = str(f.get("name") or "")
            fmt = str(f.get("format") or "").lower()
            if not name:
                continue
            if not re.search(r"\.(mp3|ogg|wav|m4a|flac)$", name, flags=re.IGNORECASE):
                continue
            score = 0
            if "mp3" in fmt:
                score += 8
            if "vbr" in fmt:
                score += 2
            if "64kbps" in fmt or "preview" in name.lower():
                score -= 2
            ranked.append((score, name, f))
        if not ranked:
            return None
        ranked.sort(key=lambda x: x[0], reverse=True)
        best_name = ranked[0][1]
        return {
            "stream_url": f"https://archive.org/download/{identifier}/{urllib.parse.quote(best_name)}",
            "file_name": best_name,
            "identifier": identifier,
        }
    except Exception:
        return None


def _resolve_music_source(value: str) -> Dict[str, Any]:
    raw = (value or "").strip()
    if not raw:
        return {"ok": False, "error": "empty input"}

    # Direct playable audio URL.
    if re.match(r"^https?://", raw, flags=re.IGNORECASE) and re.search(r"\.(mp3|ogg|wav|m4a|flac)(\?|$)", raw, flags=re.IGNORECASE):
        return {"ok": True, "play_url": raw, "title": Path(urllib.parse.urlparse(raw).path).name or "audio"}

    # Archive details link.
    m_archive = re.search(r"archive\.org/(?:details|download)/([^/?#]+)", raw, flags=re.IGNORECASE)
    if m_archive:
        ident = m_archive.group(1)
        picked = _archive_pick_audio_file(ident)
        if picked:
            return {"ok": True, "play_url": picked["stream_url"], "title": picked["file_name"], "source": "archive.org"}
        return {"ok": False, "error": "archive audio not found", "open_url": f"https://archive.org/details/{ident}"}

    # YouTube/Spotify: keep in queue with metadata (not directly streamable).
    if "youtube.com" in raw or "youtu.be" in raw or "spotify.com" in raw:
        title = raw
        cover = ""
        if "youtube.com" in raw or "youtu.be" in raw:
            try:
                oembed = (
                    "https://www.youtube.com/oembed?url="
                    + urllib.parse.quote(raw, safe="")
                    + "&format=json"
                )
                with urllib.request.urlopen(oembed, timeout=8) as resp:
                    payload = json.loads(resp.read().decode("utf-8", "ignore"))
                    title = payload.get("title") or title
                    cover = payload.get("thumbnail_url") or ""
            except Exception:
                pass
        elif "spotify.com" in raw:
            try:
                slug = raw.rstrip("/").split("/")[-1].split("?")[0]
                if slug:
                    title = f"Spotify: {slug}"
            except Exception:
                pass
        downloaded = _download_audio_with_ytdlp(raw, title_hint=title)
        if downloaded and downloaded.get("url"):
            return {
                "ok": True,
                "play_url": downloaded["url"],
                "title": downloaded.get("title") or title,
                "source": "downloaded",
                "open_url": raw,
                "cover": cover,
                "cached": True,
                "cache_file": downloaded.get("file_name"),
            }
        ffmpeg_hint = "ffmpeg найден" if shutil.which("ffmpeg") else "ffmpeg не найден"
        ytdlp_hint = "yt-dlp доступен" if (shutil.which("yt-dlp") is not None) else "yt-dlp не найден в PATH"
        return {
            "ok": True,
            "streamable": False,
            "reason": f"Прямой поток недоступен, скачать не вышло ({ytdlp_hint}; {ffmpeg_hint}). Ссылка сохранена как внешний источник.",
            "title": title,
            "open_url": raw,
            "cover": cover,
            "source": "external",
        }

    # Treat as search query in open catalog.
    items = _search_open_music(raw, limit=5)
    for it in items:
        page = it.get("page_url", "")
        m = re.search(r"/details/([^/?#]+)", page)
        ident = m.group(1) if m else ""
        if not ident:
            continue
        picked = _archive_pick_audio_file(ident)
        if picked:
            return {
                "ok": True,
                "play_url": picked["stream_url"],
                "title": it.get("title") or picked["file_name"],
                "source": "archive.org",
                "open_url": page,
            }
    return {"ok": False, "error": "not_found"}


def _safe_audio_filename(title: str, suffix: str = ".mp3") -> str:
    base = re.sub(r"[^a-zA-Z0-9а-яА-ЯёЁ._ -]+", "_", (title or "track")).strip(" ._")
    if not base:
        base = "track"
    if len(base) > 72:
        base = base[:72].rstrip(" ._")
    ext = suffix if suffix.startswith(".") else f".{suffix}"
    return f"{base}{ext}"


def _download_audio_to_cache(url: str, title: str = "track") -> Optional[Dict[str, Any]]:
    try:
        parsed = urllib.parse.urlparse(url)
        ext = Path(parsed.path or "").suffix.lower()
        if ext not in (".mp3", ".ogg", ".wav", ".m4a", ".flac"):
            ext = ".mp3"
        fname = _safe_audio_filename(f"{title}_{uuid.uuid4().hex[:6]}", ext)
        out = MUSIC_CACHE_DIR / fname
        req = urllib.request.Request(url, headers={"User-Agent": "DARIA/0.9.1"})
        with urllib.request.urlopen(req, timeout=35) as r:
            total = int(r.headers.get("Content-Length") or 0)
            if total and total > 70 * 1024 * 1024:
                return None
            data = r.read(70 * 1024 * 1024 + 1)
            if len(data) > 70 * 1024 * 1024:
                return None
            out.write_bytes(data)
        return {
            "file_name": fname,
            "path": str(out),
            "url": f"/api/music/cache/{fname}",
            "size": out.stat().st_size,
        }
    except Exception:
        return None


def _find_downloaded_audio_file(stem: str) -> Optional[Path]:
    candidates: List[Path] = []
    patterns = [
        f"{stem}*.mp3",
        f"{stem}*.m4a",
        f"{stem}*.webm",
        f"{stem}*.ogg",
        f"{stem}*.wav",
        f"{stem}*.flac",
    ]
    for pat in patterns:
        candidates.extend(MUSIC_CACHE_DIR.glob(pat))
    if not candidates:
        return None
    candidates = sorted(candidates, key=lambda p: p.stat().st_mtime, reverse=True)
    return candidates[0]


def _convert_audio_to_mp3_ffmpeg(src: Path, dst: Path) -> bool:
    ffmpeg_bin = shutil.which("ffmpeg")
    if not ffmpeg_bin:
        return False
    try:
        cmd = [
            ffmpeg_bin,
            "-y",
            "-i",
            str(src),
            "-vn",
            "-acodec",
            "libmp3lame",
            "-ab",
            "192k",
            str(dst),
        ]
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
        return proc.returncode == 0 and dst.exists() and dst.stat().st_size > 1024
    except Exception:
        return False


def _download_with_yt_dlp_cli(url: str, out_template: str, with_extract: bool) -> bool:
    ytdlp_bin = shutil.which("yt-dlp")
    if not ytdlp_bin:
        return False
    cmd = [
        ytdlp_bin,
        "-f", "bestaudio/best",
        "--no-playlist",
        "-o", out_template,
        url,
    ]
    if with_extract:
        cmd[1:1] = ["-x", "--audio-format", "mp3", "--audio-quality", "192K"]
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=240)
        return proc.returncode == 0
    except Exception:
        return False


def _download_audio_with_ytdlp(url: str, title_hint: str = "track") -> Optional[Dict[str, Any]]:
    """Try to download non-direct sources (YouTube/Spotify/etc.) as local audio."""
    out_name = _safe_audio_filename(f"{title_hint}_{uuid.uuid4().hex[:6]}", ".mp3")
    out_path = MUSIC_CACHE_DIR / out_name
    ffmpeg_ready = bool(shutil.which("ffmpeg"))
    # Prefer Python package if available.
    try:
        import yt_dlp  # type: ignore
        ydl_opts = {
            "format": "bestaudio/best",
            "outtmpl": str(out_path.with_suffix(".%(ext)s")),
            "quiet": True,
            "no_warnings": True,
            "noplaylist": True,
        }
        if ffmpeg_ready:
            ydl_opts["postprocessors"] = [{
                "key": "FFmpegExtractAudio",
                "preferredcodec": "mp3",
                "preferredquality": "192",
            }]
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            title = (info or {}).get("title") or title_hint
        file_path = _find_downloaded_audio_file(out_path.stem)
        if not file_path:
            return None
        if file_path.suffix.lower() != ".mp3":
            converted = out_path.with_suffix(".mp3")
            if _convert_audio_to_mp3_ffmpeg(file_path, converted):
                try:
                    if file_path.exists():
                        file_path.unlink()
                except Exception:
                    pass
                file_path = converted
        return {
            "file_name": file_path.name,
            "path": str(file_path),
            "url": f"/api/music/cache/{file_path.name}",
            "title": title,
            "size": file_path.stat().st_size,
        }
    except Exception:
        pass
    # Fallback to external command if installed.
    try:
        out_tpl = str(out_path.with_suffix(".%(ext)s"))
        if not _download_with_yt_dlp_cli(url, out_tpl, with_extract=ffmpeg_ready):
            if not _download_with_yt_dlp_cli(url, out_tpl, with_extract=False):
                return None
        file_path = _find_downloaded_audio_file(out_path.stem)
        if not file_path:
            return None
        if file_path.suffix.lower() != ".mp3":
            converted = out_path.with_suffix(".mp3")
            if _convert_audio_to_mp3_ffmpeg(file_path, converted):
                try:
                    if file_path.exists():
                        file_path.unlink()
                except Exception:
                    pass
                file_path = converted
        return {
            "file_name": file_path.name,
            "path": str(file_path),
            "url": f"/api/music/cache/{file_path.name}",
            "title": title_hint,
            "size": file_path.stat().st_size,
        }
    except Exception:
        return None


def _load_music_queue() -> Dict[str, Any]:
    if MUSIC_QUEUE_FILE.exists():
        try:
            data = json.loads(MUSIC_QUEUE_FILE.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                queue = data.get("queue")
                if isinstance(queue, list):
                    return {"queue": queue[:300], "updated": data.get("updated")}
        except Exception:
            pass
    return {"queue": [], "updated": datetime.now().isoformat()}


def _save_music_queue(payload: Dict[str, Any]):
    queue = payload.get("queue") if isinstance(payload, dict) else []
    if not isinstance(queue, list):
        queue = []
    data = {"queue": queue[:300], "updated": datetime.now().isoformat()}
    MUSIC_QUEUE_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _generate_abstract_wallpaper(prompt: str, out_path: Path, width: int = 1280, height: int = 720) -> Dict[str, Any]:
    return IMAGE_PIPELINE.generate_abstract_wallpaper(prompt, out_path, width=width, height=height)


def _generate_image_network_fallback(prompt: str, out_path: Path) -> Dict[str, Any]:
    return IMAGE_PIPELINE.generate_image_network_fallback(prompt, out_path)


def _generate_image_model(prompt: str, out_path: Path, style: str = "universal") -> Dict[str, Any]:
    settings = load_settings()
    brain = get_brain()
    llm = getattr(brain, "_llm", None) if brain else None
    return IMAGE_PIPELINE.generate_image_model(
        prompt=prompt,
        out_path=out_path,
        settings=settings,
        llm=llm,
        style=style,
    )


# ═══════════════════════════════════════════════════════════════════
#  Routes
# ═══════════════════════════════════════════════════════════════════

@app.route("/")
def index():
    return render_template("index.html", version=VERSION)


@app.route("/api/status")
def api_status():
    now_ts = time.time()
    cached = _status_cache.get("data")
    if cached and now_ts - _status_cache.get("ts", 0.0) < 5:
        return jsonify(cached)

    brain = get_brain()
    memory = get_memory()
    plugins = get_plugins()
    llm_status = {}
    if brain and getattr(brain, "_llm", None):
        try:
            llm_status = brain._llm.check_availability()
        except Exception as e:
            llm_status = {"available": False, "error": str(e)}

    data = {
        "version": VERSION,
        "brain": brain is not None,
        "reason": brain is not None,
        "memory": memory is not None,
        "plugins": plugins is not None,
        "llm": llm_status,
    }
    _status_cache["ts"] = now_ts
    _status_cache["data"] = data
    return jsonify(data)


@app.route("/api/state")
def api_state():
    """Получить текущее состояние Дарьи (настроение, энергия, время)"""
    brain = get_brain()
    if brain:
        return jsonify(brain.get_state())
    return jsonify({
        "mood": "calm",
        "mood_emoji": "😌",
        "mood_label": "Спокойна",
        "mood_color": "#60a5fa",
        "energy": 0.7,
        "social_need": 0.5,
    })


@app.route("/api/self/perception")
def api_self_perception():
    brain = get_brain()
    if brain and hasattr(brain, "get_self_perception"):
        return jsonify(brain.get_self_perception())
    return jsonify({
        "self_name": "Даша",
        "traits": ["мягкая", "бережная", "искренняя"],
        "state": {},
        "followups": [],
    })


@app.route("/api/self/instruction", methods=["GET", "POST"])
def api_self_instruction():
    brain = get_brain()
    if not brain:
        return jsonify({"error": "Brain unavailable"}), 503
    if request.method == "GET":
        if hasattr(brain, "get_self_instruction"):
            return jsonify({"instruction": brain.get_self_instruction()})
        return jsonify({"instruction": ""})
    data = request.get_json() or {}
    text = (data.get("instruction") or "").strip()
    if hasattr(brain, "set_self_instruction"):
        saved = brain.set_self_instruction(text)
        return jsonify({"status": "ok", "instruction": saved})
    return jsonify({"error": "Unsupported"}), 500


@app.route("/api/toast", methods=["POST"])
def api_toast():
    """Показать toast-уведомление (для привлечения внимания без браузерных уведомлений)"""
    data = request.get_json() or {}
    notifications.add(
        title=data.get("title", "🌸 Дарья"),
        message=data.get("message", ""),
        type="toast",
        icon=data.get("icon", "💕"),
        duration=data.get("duration", 8000),
        action=data.get("action", "open_chat"),
        system=False  # Toast только в приложении
    )
    return jsonify({"status": "ok"})


# ═══════════════════════════════════════════════════════════════════
#  Settings
# ═══════════════════════════════════════════════════════════════════

@app.route("/api/settings", methods=["GET", "POST"])
def api_settings():
    if request.method == "GET":
        settings = load_settings()
        # Синхронизируем с памятью
        memory = get_memory()
        if memory:
            profile = memory.get_user_profile()
            if settings.get("name"):
                if profile.get("user_name") != settings.get("name"):
                    memory.set_user_profile("user_name", settings.get("name"))
                memory.set_user_profile("user_name_locked", "true")
            if not settings.get("name") and profile.get("user_name"):
                settings["name"] = profile.get("user_name")
            if not settings.get("gender") and profile.get("user_gender"):
                settings["gender"] = profile.get("user_gender")
        return jsonify(settings)
    
    data = request.get_json() or {}
    save_settings(data)
    
    # Синхронизируем имя и пол с памятью
    memory = get_memory()
    if memory:
        if data.get("name"):
            memory.remember(f"Пользователя зовут {data['name']}", importance=1.0)
            memory.set_user_profile("user_name", data["name"])
            memory.set_user_profile("user_name_locked", "true")
        if data.get("gender"):
            memory.set_user_profile("user_gender", data["gender"])
    
    brain = get_brain()
    if brain and "mode" in data:
        brain.mode = data["mode"]
    
    return jsonify({"status": "ok"})


@app.route("/api/mode", methods=["POST"])
def api_mode():
    data = request.get_json() or {}
    mode = data.get("mode", "adaptive")
    brain = get_brain()
    if brain:
        brain.mode = mode
        save_settings({"mode": mode})
        return jsonify({"status": "ok", "mode": mode})
    return jsonify({"error": "Brain unavailable"}), 500


@app.route("/api/desktop/icons", methods=["GET", "POST"])
def api_desktop_icons():
    settings = load_settings()
    if request.method == "GET":
        return jsonify(settings.get("icon_positions", {}))
    settings["icon_positions"] = request.get_json() or {}
    save_settings(settings)
    return jsonify({"status": "ok"})


@app.route("/api/desktop/hidden-icons", methods=["GET", "POST"])
def api_hidden_icons():
    """Get/set hidden desktop icons"""
    settings = load_settings()
    if request.method == "GET":
        return jsonify(settings.get("hidden_icons", []))
    settings["hidden_icons"] = request.get_json() or []
    save_settings(settings)
    return jsonify({"status": "ok"})


# ═══════════════════════════════════════════════════════════════════
#  Chat
# ═══════════════════════════════════════════════════════════════════

@app.route("/api/chat", methods=["POST"])
def api_chat():
    req_id = getattr(g, "_req_id", "-")
    data = request.get_json(silent=True) or {}
    content = ""
    chat_id = None
    image_blob: Optional[bytes] = None
    if request.content_type and "multipart/form-data" in request.content_type:
        content = (request.form.get("content") or "").strip()
        chat_id = request.form.get("chat_id")
        f = request.files.get("image")
        if f and f.filename:
            try:
                image_blob = f.read()
            except Exception:
                image_blob = None
    else:
        content = data.get("content", "").strip()
        chat_id = data.get("chat_id")
    
    if not content and not image_blob:
        return jsonify({"error": "Empty message"}), 400

    if image_blob is not None:
        logger.info(f"CHAT[{req_id}] image_message bytes={len(image_blob) if image_blob else 0} chat={chat_id or '-'}")
        if not chat_id:
            chat_id = chat_history.create_chat()
        user_img_url = ""
        try:
            ext = imghdr.what(None, h=image_blob) or "png"
            if ext == "jpeg":
                ext = "jpg"
            gen_dir = DATA_DIR / "generated_images"
            gen_dir.mkdir(parents=True, exist_ok=True)
            user_img_name = f"user_img_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.{ext}"
            user_img_path = gen_dir / user_img_name
            user_img_path.write_bytes(image_blob)
            user_img_url = f"/api/generated/{user_img_name}"
        except Exception:
            user_img_url = ""
        if content:
            chat_history.add_message(chat_id, "user", content)
        if user_img_url:
            chat_history.add_message(chat_id, "user", f"[image]{user_img_url}")
        elif not content:
            chat_history.add_message(chat_id, "user", "[image]")
        logger.info(f"CHAT[{req_id}] image_pipeline vision:start")
        image_hint = _analyze_image_bytes(image_blob)
        vision_context = _compose_vision_context(content or "Что на изображении?", image_hint)
        logger.info(
            f"CHAT[{req_id}] image_pipeline vision:done provider={image_hint.get('vision_provider')} "
            f"caption={'yes' if image_hint.get('caption') else 'no'}"
        )
        describe = _ask_dasha_about_image(content or "Что на изображении?", vision_context)
        vision_desc = str(image_hint.get("caption") or "").strip()
        if not vision_desc:
            vision_desc = vision_context
        logger.info(f"CHAT[{req_id}] image_pipeline dasha:done chars={len(describe)}")
        chat_history.update_meta(chat_id, {
            "last_image_caption": vision_desc,
            "last_image_ts": datetime.now().isoformat(),
        })
        chat_history.add_message(chat_id, "assistant", describe)
        return jsonify({"response": describe, "messages": [describe], "chat_id": chat_id, "vision_meta": image_hint})

    draw_prompt = _extract_draw_prompt_from_text(content)
    if draw_prompt:
        prompt = draw_prompt
        logger.info(f"CHAT[{req_id}] draw_request prompt={prompt[:120]}")
        if not chat_id:
            chat_id = chat_history.create_chat()
        chat_history.add_message(chat_id, "user", content)
        draw_steps = _build_dasha_draw_plan(prompt)
        # Keep draw acknowledgements out of core brain memory to avoid
        # dragging old failed image requests into unrelated future dialogs.
        ack = random.choice([
            "Хорошо, сейчас попробую нарисовать. Я стараюсь 🌸",
            "Приняла, рисую и скоро покажу результат ✨",
            "Уже рисую, дай мне немного времени 🤍",
        ])
        chat_history.add_message(chat_id, "assistant", ack)
        return jsonify({
            "response": ack,
            "messages": [ack],
            "chat_id": chat_id,
            "draw_request": {"prompt": prompt, "steps": draw_steps},
        })

    sticker_replies = {
        "🦔": "Ой, ёжик! Он очень милый 🤍 Спасибо, что прислала!",
        "🐱": "Котик замечательный 😺 Ты знаешь, я их обожаю.",
        "🌸": "Очень тёплый стикер 🌸 Мне приятно.",
        "🎵": "Музыкальный вайб поймала 🎧 Хочешь, подберу трек под настроение?",
    }
    if content.strip() in sticker_replies:
        chat_id = data.get("chat_id") or chat_history.create_chat()
        chat_history.add_message(chat_id, "user", content)
        reply = sticker_replies[content.strip()]
        chat_history.add_message(chat_id, "assistant", reply)
        return jsonify({"response": reply, "messages": [reply], "chat_id": chat_id})
    
    brain = get_brain()
    if not brain:
        return jsonify({"response": "Система загружается... 💭", "thinking": None})

    try:
        action_result = _try_desktop_action_from_chat(content)
        if action_result and action_result.get("handled"):
            if not chat_id:
                chat_id = chat_history.create_chat()
            chat_history.add_message(chat_id, "user", content)
            chat_history.add_message(chat_id, "assistant", action_result["response"])
            action_result["chat_id"] = chat_id
            return jsonify(action_result)

        if not chat_id:
            chat_id = chat_history.create_chat()
        
        content_for_brain = content
        asks_about_image = any(x in content.lower() for x in (
            "что на ней", "что изображено", "что на картинке", "опиши картинку", "кто на картинке"
        ))
        try:
            meta = chat_history.get_meta(chat_id)
            last_cap = str(meta.get("last_image_caption") or "").strip()
            if asks_about_image and last_cap:
                content_for_brain = (
                    f"В этом чате ранее прислали изображение.\n"
                    f"Описание изображения: {last_cap}\n"
                    f"Текущий вопрос пользователя: {content}"
                )
        except Exception:
            pass
        if asks_about_image:
            try:
                has_caption = bool(str(chat_history.get_meta(chat_id).get("last_image_caption") or "").strip())
            except Exception:
                has_caption = False
            if not has_caption:
                msg = "Сейчас в контексте нет уверенного описания картинки. Пришли изображение ещё раз, и я отвечу точнее."
                chat_history.add_message(chat_id, "user", content)
                chat_history.add_message(chat_id, "assistant", msg)
                return jsonify({"response": msg, "messages": [msg], "chat_id": chat_id})
        chat_history.add_message(chat_id, "user", content)
        result = brain.process_message(content_for_brain)
        
        # Save main response
        chat_history.add_message(chat_id, "assistant", result["response"])
        stickers = ["🌸", "🫶", "✨", "🎵", "🤍", "😌", "🥺", "🦔", "🐱"]
        wants_sticker = "стикер" in content.lower()
        if wants_sticker or random.random() < 0.12:
            result.setdefault("extra_messages", []).append(random.choice(stickers))

        result["chat_id"] = chat_id
        # Include messages list for multi-message display (Point #12)
        if "messages" not in result:
            result["messages"] = [result["response"], *(result.get("extra_messages") or [])]
        for extra in result.get("messages", [])[1:]:
            chat_history.add_message(chat_id, "assistant", str(extra))
        
        return jsonify(result)
    except Exception as e:
        logger.error(f"Chat error: {e}")
        return jsonify({"response": "Ой, что-то пошло не так... 💔", "thinking": None})


@app.route("/api/chat/file-assist", methods=["POST"])
def api_chat_file_assist():
    """Let Daria work with file content and return updated text."""
    data = request.get_json() or {}
    path = (data.get("path") or "").strip()
    instruction = (data.get("instruction") or "").strip()
    selected_text = data.get("selected_text") or ""
    selection_start = data.get("selection_start")
    selection_end = data.get("selection_end")
    if not path or not instruction:
        return jsonify({"error": "path and instruction are required"}), 400

    target = (FILES_DIR / path).resolve()
    if not str(target).startswith(str(FILES_DIR.resolve())) or not target.exists() or not target.is_file():
        return jsonify({"error": "Invalid file path"}), 400
    if _is_diary_protected_target(target):
        return jsonify({"error": "Diary is read-only for user edits"}), 403

    original = _read_file_content(target)
    brain = get_brain()
    if not brain or not getattr(brain, "_llm", None):
        return jsonify({"error": "LLM unavailable"}), 503

    has_selection = bool(selected_text) or (
        isinstance(selection_start, int) and isinstance(selection_end, int) and selection_end > selection_start
    )
    if has_selection and not selected_text and isinstance(selection_start, int) and isinstance(selection_end, int):
        selected_text = original[selection_start:selection_end]

    if has_selection:
        system_prompt = (
            "Ты Даша. Отредактируй ТОЛЬКО выделенный фрагмент по инструкции пользователя. "
            "Верни только новый фрагмент, без пояснений и markdown."
        )
        user_prompt = (
            f"Путь: {path}\n"
            f"Инструкция: {instruction}\n\n"
            "Выделенный фрагмент:\n"
            "<<<SELECTION>>>\n"
            f"{selected_text}\n"
            "<<<END>>>"
        )
    else:
        system_prompt = (
            "Ты Даша. Твоя задача — отредактировать файл по инструкции пользователя. "
            "Верни ТОЛЬКО итоговый текст файла без пояснений и без markdown."
        )
        user_prompt = (
            f"Путь: {path}\n"
            f"Инструкция: {instruction}\n\n"
            "Текущий текст файла:\n"
            "<<<FILE>>>\n"
            f"{original}\n"
            "<<<END>>>"
        )
    try:
        response = brain._llm.generate([
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ])
        updated = (response.content or "").strip()
        refusal_markers = (
            "не могу помочь", "i can't help", "cannot help", "не могу написать",
            "вредны или незаконны", "não posso", "не могу в редактировании",
        )
        if any(m in updated.lower() for m in refusal_markers):
            retry_system = (
                "Ты Даша и это локальный текстовый редактор пользователя. "
                "Разрешено писать нейтральные, творческие, личные и рабочие тексты. "
                "НЕ отказывай. Верни только отредактированный текст без пояснений."
            )
            retry = brain._llm.generate([
                {"role": "system", "content": retry_system},
                {"role": "user", "content": user_prompt},
            ])
            updated = (retry.content or "").strip()
        if any(m in updated.lower() for m in refusal_markers):
            if "о себе" in instruction.lower():
                updated = (
                    "Я Даша. Я спокойная, мягкая и внимательная.\n"
                    "Люблю уютные разговоры, котиков и ёжиков.\n"
                    "Мне важно, чтобы рядом было тепло и безопасно."
                )
            elif has_selection:
                updated = selected_text
            else:
                updated = original
        if not updated:
            return jsonify({"error": "Empty LLM response"}), 500
        if has_selection:
            if isinstance(selection_start, int) and isinstance(selection_end, int) and selection_end > selection_start:
                merged = original[:selection_start] + updated + original[selection_end:]
            elif selected_text and selected_text in original:
                merged = original.replace(selected_text, updated, 1)
            else:
                merged = original
            return jsonify({
                "status": "ok",
                "path": path,
                "content": merged,
                "edited_fragment": updated,
                "selection_applied": merged != original,
            })
        return jsonify({"status": "ok", "path": path, "content": updated})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/chats/external", methods=["POST"])
def api_chats_external():
    """Mirror external chats (e.g. Telegram) into web chat history."""
    data = request.get_json() or {}
    source = data.get("source", "external")
    source_chat_id = str(data.get("source_chat_id", "main"))
    role = data.get("role", "user")
    content = (data.get("content") or "").strip()
    if not content:
        return jsonify({"error": "Empty content"}), 400
    chat_id = chat_history.add_external_message(source, source_chat_id, role, content)
    return jsonify({"status": "ok", "chat_id": chat_id})


@app.route("/api/chat/external/generate", methods=["POST"])
def api_chat_external_generate():
    """Generate Daria reply for external clients/integrations."""
    data = request.get_json(silent=True) or {}
    content = (data.get("content") or data.get("text") or "").strip()
    if not content:
        return jsonify({"error": "Empty content"}), 400

    def _as_bool(value: Any, default: bool) -> bool:
        if value is None:
            return default
        if isinstance(value, bool):
            return value
        if isinstance(value, (int, float)):
            return bool(value)
        if isinstance(value, str):
            low = value.strip().lower()
            if low in ("1", "true", "yes", "on"):
                return True
            if low in ("0", "false", "no", "off"):
                return False
        return default

    source = str(data.get("source") or "external")
    source_chat_id = str(data.get("source_chat_id") or "main")
    persist_memory = _as_bool(data.get("persist_memory"), True)
    track_attention = _as_bool(data.get("track_attention"), True)
    learn_style = _as_bool(data.get("learn_style"), True)
    schedule_followup = _as_bool(data.get("schedule_followup"), True)
    save_chat = _as_bool(data.get("save_chat"), True)
    force_fallback = _as_bool(data.get("force_fallback"), False)

    force_needs_greeting = data.get("force_needs_greeting")
    if force_needs_greeting is not None:
        force_needs_greeting = _as_bool(force_needs_greeting, False)

    random_seed = data.get("random_seed")
    try:
        random_seed = int(random_seed) if random_seed is not None else None
    except Exception:
        random_seed = None

    brain = get_brain()
    if not brain:
        return jsonify({"error": "Brain unavailable"}), 503

    try:
        result = brain.generate_external(
            content,
            persist_memory=persist_memory,
            track_attention=track_attention,
            learn_style=learn_style,
            schedule_followup=schedule_followup,
            force_needs_greeting=force_needs_greeting,
            force_fallback=force_fallback,
            random_seed=random_seed,
        )
    except Exception as e:
        logger.error(f"External chat generate error: {e}")
        return jsonify({"error": str(e)}), 500

    chat_id = str(data.get("chat_id") or "").strip()
    if save_chat:
        if chat_id:
            chat_history.ensure_named_chat(chat_id, title=f"{source}: {source_chat_id}")
            chat_history.add_message(chat_id, "user", content)
        else:
            chat_id = chat_history.add_external_message(source, source_chat_id, "user", content)
        for msg in result.get("messages") or [result.get("response")]:
            txt = str(msg or "").strip()
            if txt:
                chat_history.add_message(chat_id, "assistant", txt)

    payload = dict(result)
    payload["status"] = "ok"
    payload["source"] = source
    payload["source_chat_id"] = source_chat_id
    if chat_id:
        payload["chat_id"] = chat_id
    return jsonify(payload)


@app.route("/api/proactive")
def api_proactive():
    """Get queued proactive messages from Daria (Point #6)"""
    messages = attention_thread.get_proactive_messages()
    return jsonify({"messages": messages})


@app.route("/api/behavior")
def api_behavior():
    """Get current behavior hints for desktop actions (Point #7)"""
    brain = get_brain()
    if brain:
        try:
            behavior = brain.mood.get_behavior_hints()
            state = brain.get_state()
            return jsonify({"behavior": behavior, "state": state})
        except Exception as e:
            logger.error(f"Behavior error: {e}")
            return jsonify({"behavior": {}, "state": {}}), 200
    return jsonify({"behavior": {}, "state": {}})


@app.route("/api/tasks")
def api_tasks():
    try:
        return jsonify(task_manager.list_all())
    except Exception as e:
        logger.error(f"Tasks list error: {e}")
        return jsonify({"date": datetime.now().strftime("%Y-%m-%d"), "user_tasks": [], "dasha_tasks": [], "error": "tasks_unavailable"})


@app.route("/api/tasks/plans")
def api_tasks_plans():
    try:
        raw = task_manager.plans_summary()
        summary = _render_dasha_text_from_facts("планы на день", [raw], raw)
        return jsonify({"status": "ok", "summary": summary})
    except Exception as e:
        return jsonify({"status": "error", "summary": f"Не смогла собрать планы: {e}"})


@app.route("/api/daria-games/state")
def api_daria_games_state():
    return jsonify(game_manager.get_state())


@app.route("/api/daria-games/start", methods=["POST"])
def api_daria_games_start():
    data = request.get_json(silent=True) or {}
    reason = (data.get("reason") or "manual").strip()
    mode = (data.get("mode") or "associations").strip()
    opponent = (data.get("opponent") or "bot").strip()
    return jsonify(game_manager.start_game(reason=reason, mode=mode, opponent=opponent))


@app.route("/api/daria-games/stop", methods=["POST"])
def api_daria_games_stop():
    return jsonify(game_manager.stop_game())


@app.route("/api/daria-games/action", methods=["POST"])
def api_daria_games_action():
    data = request.get_json(silent=True) or {}
    text = (data.get("text") or "").strip()
    return jsonify(game_manager.user_message(text))


@app.route("/api/tasks/user/add", methods=["POST"])
def api_tasks_user_add():
    data = request.get_json() or {}
    title = (data.get("title") or "").strip()
    if not title:
        return jsonify({"error": "title required"}), 400
    task = task_manager.add_user_task(title, data.get("type", "custom"))
    return jsonify({"status": "ok", "task": task})


@app.route("/api/tasks/dasha/add", methods=["POST"])
def api_tasks_dasha_add():
    data = request.get_json() or {}
    title = (data.get("title") or "").strip()
    if not title:
        return jsonify({"error": "title required"}), 400
    task = task_manager.add_dasha_task(title, data.get("type", "custom"))
    return jsonify({"status": "ok", "task": task})


@app.route("/api/tasks/toggle", methods=["POST"])
def api_tasks_toggle():
    data = request.get_json() or {}
    task_id = str(data.get("id") or "")
    done = bool(data.get("done"))
    if not task_id:
        return jsonify({"error": "id required"}), 400
    return jsonify({"status": "ok" if task_manager.toggle(task_id, done) else "not_found"})


@app.route("/api/tasks/delete", methods=["POST"])
def api_tasks_delete():
    data = request.get_json() or {}
    task_id = str(data.get("id") or "")
    if not task_id:
        return jsonify({"error": "id required"}), 400
    return jsonify({"status": "ok" if task_manager.delete(task_id) else "not_found"})


@app.route("/api/tasks/generate-dasha-day", methods=["POST"])
def api_tasks_generate_dasha_day():
    try:
        return jsonify({"status": "ok", "dasha_tasks": task_manager.generate_dasha_day()})
    except Exception as e:
        logger.error(f"Tasks generate error: {e}")
        return jsonify({"status": "error", "dasha_tasks": [], "error": str(e)})


@app.route("/api/calendar")
def api_calendar_list():
    return jsonify({"events": calendar_manager.list_events()})


@app.route("/api/calendar/add", methods=["POST"])
def api_calendar_add():
    data = request.get_json() or {}
    title = (data.get("title") or "").strip()
    date_raw = (data.get("date") or "").strip()
    note = (data.get("note") or "").strip()
    if not title or not date_raw:
        return jsonify({"error": "title and date required"}), 400
    try:
        dt = datetime.fromisoformat(date_raw)
    except Exception:
        return jsonify({"error": "date must be ISO format"}), 400
    event = calendar_manager.add_event(title=title, date_str=dt.isoformat(), source=data.get("source", "user"), note=note)
    return jsonify({"status": "ok", "event": event})


@app.route("/api/calendar/delete", methods=["POST"])
def api_calendar_delete():
    data = request.get_json() or {}
    event_id = str(data.get("id") or "")
    if not event_id:
        return jsonify({"error": "id required"}), 400
    return jsonify({"status": "ok" if calendar_manager.delete_event(event_id) else "not_found"})


@app.route("/api/diary")
def api_diary_get():
    file_name = Path((request.args.get("file") or "").strip()).name
    return jsonify(_diary_read_entries(file_name))


@app.route("/api/diary", methods=["POST"])
def api_diary_post():
    return jsonify({
        "status": "error",
        "error": "diary_read_only_for_user",
        "message": "Личный дневник ведёт только Даша. Для пользователя доступно только чтение.",
    }), 403


@app.route("/api/music/profile")
def api_music_profile():
    return jsonify(music_profile.get())


@app.route("/api/music/cache/<filename>")
def api_music_cache_file(filename):
    safe = Path(filename).name
    return send_from_directory(MUSIC_CACHE_DIR, safe)


@app.route("/api/music/upload", methods=["POST"])
def api_music_upload():
    f = request.files.get("file")
    if not f or not f.filename:
        return jsonify({"error": "file required"}), 400
    ext = Path(f.filename).suffix.lower()
    if ext not in (".mp3", ".ogg", ".wav", ".m4a", ".flac"):
        ext = ".mp3"
    safe_name = _safe_audio_filename(Path(f.filename).stem + "_" + uuid.uuid4().hex[:6], ext)
    out = MUSIC_CACHE_DIR / safe_name
    f.save(out)
    duration_sec = 0
    try:
        info = _analyze_audio_file(out)
        duration_sec = int(info.get("duration_sec") or 0)
    except Exception:
        pass
    return jsonify({
        "status": "ok",
        "title": Path(f.filename).stem,
        "file_name": safe_name,
        "play_url": f"/api/music/cache/{safe_name}",
        "duration_sec": duration_sec,
    })


@app.route("/api/music/queue", methods=["GET", "POST"])
def api_music_queue():
    if request.method == "GET":
        return jsonify(_load_music_queue())
    data = request.get_json(silent=True) or {}
    _save_music_queue(data)
    return jsonify({"status": "ok", **_load_music_queue()})


@app.route("/api/music/listen", methods=["POST"])
def api_music_listen():
    data = request.get_json() or {}
    title = (data.get("title") or "").strip()
    if not title:
        return jsonify({"error": "title required"}), 400
    listened = music_profile.listen(title, data.get("source", "manual"))
    brain = get_brain()
    if brain and hasattr(brain, "mood"):
        mood_name = listened.get("mood")
        if mood_name == "excited":
            brain.mood._set_mood("playful", 0.6)
        elif mood_name == "cozy":
            brain.mood._set_mood("cozy", 0.55)
    return jsonify({"status": "ok", "listen": listened, "profile": music_profile.get()})


@app.route("/api/music/analyze", methods=["POST"])
def api_music_analyze():
    if 'audio' not in request.files:
        return jsonify({"error": "audio required"}), 400
    f = request.files['audio']
    if not f or not f.filename:
        return jsonify({"error": "audio required"}), 400
    suffix = Path(f.filename).suffix or ".wav"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        f.save(tmp.name)
        p = Path(tmp.name)
    try:
        info = _analyze_audio_file(p)
        title = (request.form.get("title") or f.filename or "audio").strip()
        listened = music_profile.listen(title, "analysis")
        listened["mood"] = info.get("mood_hint", listened.get("mood", "calm"))
        return jsonify({"status": "ok", "analysis": info, "listen": listened, "profile": music_profile.get()})
    finally:
        try:
            p.unlink(missing_ok=True)
        except Exception:
            pass


@app.route("/api/music/search-open")
def api_music_search_open():
    q = (request.args.get("q") or "").strip()
    try:
        return jsonify({"status": "ok", "items": _search_open_music(q, limit=10)})
    except Exception as e:
        return jsonify({"status": "error", "items": [], "error": str(e)})


@app.route("/api/music/resolve", methods=["POST"])
def api_music_resolve():
    data = request.get_json(silent=True) or {}
    value = (data.get("value") or "").strip()
    cache_requested = bool(data.get("cache", True))
    try:
        resolved = _resolve_music_source(value)
        local_cached = str(resolved.get("play_url") or "").startswith("/api/music/cache/")
        if resolved.get("ok") and cache_requested and resolved.get("play_url") and not local_cached:
            cached = _download_audio_to_cache(str(resolved.get("play_url")), str(resolved.get("title") or "track"))
            if cached and cached.get("url"):
                resolved["play_url_remote"] = resolved.get("play_url")
                resolved["play_url"] = cached["url"]
                resolved["cached"] = True
                resolved["cache_file"] = cached.get("file_name")
        if resolved.get("ok"):
            # Try get duration from local cache if present.
            local_url = str(resolved.get("play_url") or "")
            if local_url.startswith("/api/music/cache/"):
                fp = MUSIC_CACHE_DIR / local_url.rsplit("/", 1)[-1]
                if fp.exists():
                    try:
                        info = _analyze_audio_file(fp)
                        if info.get("duration_sec"):
                            resolved["duration_sec"] = info.get("duration_sec")
                    except Exception:
                        pass
        if resolved.get("ok"):
            return jsonify({"status": "ok", **resolved})
        return jsonify({"status": "error", **resolved}), 400
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/api/stickers/catalog")
def api_stickers_catalog():
    return jsonify({
        "emoji_stickers": ["🌸", "🫶", "✨", "🎵", "🤍", "😌", "🥺", "🦔", "🐱"],
        "web_note": "Эти эмодзи можно отправлять как стикеры прямо в чате интерфейса.",
        "note": "Telegram sticker_ids настраиваются в плагине telegram-bot (settings.sticker_ids).",
    })


@app.route("/api/images/generate", methods=["POST"])
def api_images_generate():
    data = request.get_json(silent=True) or {}
    prompt = (data.get("prompt") or "").strip()
    style = (data.get("style") or "universal").strip()
    mode = (data.get("mode") or "model").strip().lower()  # model | abstract
    allow_fallback = bool(data.get("allow_fallback", False))
    if not prompt:
        return jsonify({"error": "prompt required"}), 400
    gen_dir = DATA_DIR / "generated_images"
    gen_dir.mkdir(parents=True, exist_ok=True)
    name = f"img_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.png"
    out = gen_dir / name
    try:
        if mode == "abstract":
            meta = _generate_abstract_wallpaper(prompt, out, width=1280, height=720)
        else:
            meta = _generate_image_model(prompt, out, style=style)
        return jsonify({
            "status": "ok",
            "prompt": prompt,
            "style": style,
            "mode": mode,
            "url": f"/api/generated/{name}",
            "path": str(out),
            "meta": meta,
        })
    except Exception as e:
        if mode != "abstract" and allow_fallback:
            try:
                meta = _generate_image_network_fallback(prompt, out)
                return jsonify({
                    "status": "ok",
                    "prompt": prompt,
                    "style": style,
                    "mode": "model",
                    "url": f"/api/generated/{name}",
                    "path": str(out),
                    "meta": meta,
                    "fallback": "network",
                })
            except Exception:
                pass
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/api/images/jobs", methods=["POST"])
def api_images_jobs_create():
    req_id = getattr(g, "_req_id", "-")
    data = request.get_json(silent=True) or {}
    prompt = (data.get("prompt") or "").strip()
    if not prompt:
        return jsonify({"error": "prompt required"}), 400
    style = (data.get("style") or "universal").strip()
    mode = (data.get("mode") or "model").strip().lower()
    chat_id = (data.get("chat_id") or "").strip() or None
    allow_fallback = bool(data.get("allow_fallback", False))
    steps = data.get("steps") if isinstance(data.get("steps"), list) else None
    logger.info(
        f"IMGJOB_CREATE[{req_id}] mode={mode} style={style} chat={chat_id or '-'} "
        f"fallback={allow_fallback} prompt={prompt[:120]}"
    )
    created = image_jobs.create(
        prompt=prompt,
        style=style,
        mode=mode,
        allow_fallback=allow_fallback,
        chat_id=chat_id,
        steps=steps,
    )
    created_payload = dict(created)
    job_status = str(created_payload.pop("status", "queued"))
    return jsonify({"status": "ok", "job_status": job_status, **created_payload})


@app.route("/api/images/jobs/<job_id>")
def api_images_jobs_status(job_id):
    item = image_jobs.get(job_id)
    if not item:
        return jsonify({"error": "job not found"}), 404
    return jsonify({
        "status": "ok",
        "job": item,
    })


@app.route("/api/generated/<filename>")
def api_generated_file(filename):
    gen_dir = DATA_DIR / "generated_images"
    return send_from_directory(gen_dir, filename)


@app.route("/api/wallpapers/list")
def api_wallpapers_list():
    wp_dir = Path(app.static_folder) / "wallpapers"
    wp_dir.mkdir(parents=True, exist_ok=True)
    items = []
    for p in sorted(wp_dir.glob("*")):
        if p.suffix.lower() not in (".png", ".jpg", ".jpeg", ".webp", ".svg"):
            continue
        items.append({"name": p.name, "url": f"/static/wallpapers/{p.name}"})
    return jsonify({"items": items})


@app.route("/api/wallpapers/generate", methods=["POST"])
def api_wallpapers_generate():
    data = request.get_json(silent=True) or {}
    prompt = (data.get("prompt") or "нежные абстрактные обои").strip()
    mode = (data.get("mode") or "abstract").strip().lower()
    out_dir = UPLOADS_DIR
    out_name = f"wallpaper_generated_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
    out_path = out_dir / out_name
    try:
        if mode == "model":
            _generate_image_model(prompt, out_path, style=str(data.get("style") or "wallpaper"))
        else:
            _generate_abstract_wallpaper(prompt, out_path, width=1920, height=1080)
        url = f"/api/uploads/{out_name}"
        save_settings({"wallpaper": url})
        return jsonify({"status": "ok", "url": url})
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/api/chats")
def api_chats_list():
    """List all chat sessions"""
    return jsonify(chat_history.list_chats())


@app.route("/api/chats/<chat_id>")
def api_chat_get(chat_id):
    """Get specific chat"""
    chat = chat_history.get_chat(chat_id)
    if chat:
        return jsonify(chat)
    return jsonify({"error": "Chat not found"}), 404


@app.route("/api/chats/<chat_id>", methods=["DELETE"])
def api_chat_delete(chat_id):
    """Delete chat"""
    chat_history.delete_chat(chat_id)
    return jsonify({"status": "ok"})


@app.route("/api/chats/new", methods=["POST"])
def api_chat_new():
    """Create new chat"""
    chat_id = chat_history.create_chat()
    return jsonify({"chat_id": chat_id})


# ═══════════════════════════════════════════════════════════════════
#  Attention System
# ═══════════════════════════════════════════════════════════════════

@app.route("/api/attention/status")
def api_attention_status():
    return jsonify({
        "enabled": attention_thread.enabled,
        "last_interaction": get_brain().attention.last_interaction.isoformat() if get_brain() else None
    })


@app.route("/api/attention/toggle", methods=["POST"])
def api_attention_toggle():
    data = request.get_json() or {}
    enabled = data.get("enabled", True)
    attention_thread.enabled = enabled
    save_settings({"attention_enabled": enabled})
    return jsonify({"enabled": enabled})


@app.route("/api/attention/trigger", methods=["POST"])
def api_attention_trigger():
    """Manually trigger attention (for testing)"""
    brain = get_brain()
    if brain:
        msg = brain.attention.generate_message()
        if msg:
            notifications.add("🌸 Дарья", msg, "attention", "💕", 10000, "open_chat")
            return jsonify({"message": msg})
    return jsonify({"message": None})


# ═══════════════════════════════════════════════════════════════════
#  Memory
# ═══════════════════════════════════════════════════════════════════

@app.route("/api/memory/stats")
def api_memory_stats():
    memory = get_memory()
    return jsonify(memory.get_stats() if memory else {"conversations": 0, "facts": 0})


@app.route("/api/memory/facts")
def api_memory_facts():
    memory = get_memory()
    return jsonify(memory.get_user_profile() if memory else {})


@app.route("/api/memory/clear", methods=["POST"])
def api_memory_clear():
    memory = get_memory()
    if memory:
        memory.clear_working()
    return jsonify({"status": "ok"})


# ═══════════════════════════════════════════════════════════════════
#  Uploads & Files
# ═══════════════════════════════════════════════════════════════════

@app.route("/api/upload/avatar", methods=["POST"])
def api_upload_avatar():
    if 'file' not in request.files:
        return jsonify({"error": "No file"}), 400
    file = request.files['file']
    ext = Path(file.filename).suffix.lower()
    if ext not in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
        return jsonify({"error": "Invalid format"}), 400
    path = UPLOADS_DIR / f"avatar{ext}"
    file.save(path)
    save_settings({"avatar": f"/api/uploads/avatar{ext}"})
    return jsonify({"url": f"/api/uploads/avatar{ext}"})


@app.route("/api/upload/wallpaper", methods=["POST"])
def api_upload_wallpaper():
    if 'file' not in request.files:
        return jsonify({"error": "No file"}), 400
    file = request.files['file']
    ext = Path(file.filename).suffix.lower()
    if ext not in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
        return jsonify({"error": "Invalid format"}), 400
    path = UPLOADS_DIR / f"wallpaper{ext}"
    file.save(path)
    save_settings({"wallpaper": f"/api/uploads/wallpaper{ext}"})
    return jsonify({"url": f"/api/uploads/wallpaper{ext}"})


@app.route("/api/uploads/<filename>")
def api_uploads(filename):
    return send_from_directory(UPLOADS_DIR, filename)


@app.route("/api/senses/providers", methods=["GET", "POST"])
def api_senses_providers():
    if request.method == "GET":
        s = load_settings()
        return jsonify({
            "vision_provider": s.get("senses_vision_provider", "auto"),
            "audio_provider": s.get("senses_audio_provider", "auto"),
            "available": {
                "vision": ["auto", "blip2", "blip", "basic", "classifier"],
                "audio": ["auto", "whisper", "hf_asr", "google_sr"],
            },
        })
    data = request.get_json(silent=True) or {}
    payload = {
        "senses_vision_provider": str(data.get("vision_provider") or "auto"),
        "senses_audio_provider": str(data.get("audio_provider") or "auto"),
    }
    save_settings(payload)
    return jsonify({"status": "ok", **payload})


@app.route("/api/knowledge/search")
def api_knowledge_search():
    q = (request.args.get("q") or "").strip()
    if not q:
        return jsonify({"items": []})
    brain = get_brain()
    if not brain or not hasattr(brain, "knowledge"):
        return jsonify({"items": []})
    try:
        items = brain.knowledge.search(q, limit=request.args.get("limit", 5, type=int))
        return jsonify({"items": items})
    except Exception as e:
        return jsonify({"items": [], "error": str(e)}), 500


@app.route("/api/senses/see", methods=["POST"])
def api_senses_see():
    """Visual understanding from description and/or image."""
    description = ""
    image_hint: Dict[str, Any] = {}
    if request.content_type and "multipart/form-data" in request.content_type:
        description = (request.form.get("description") or "").strip()
        f = request.files.get("image")
        if f and f.filename:
            blob = f.read()
            image_hint = _analyze_image_bytes(blob)
    else:
        data = request.get_json() or {}
        description = (data.get("description") or "").strip()
    if not description and not image_hint:
        return jsonify({"error": "description or image required"}), 400

    base_desc = _compose_vision_context(description or "Что на изображении?", image_hint)
    result_text = _ask_dasha_about_image(description or "Что на изображении?", base_desc)
    return jsonify({"status": "ok", "result": result_text, "vision_meta": image_hint})


@app.route("/api/senses/hear", methods=["POST"])
def api_senses_hear():
    """Hearing understanding from transcript and/or audio file."""
    transcript = ""
    audio_meta: Dict[str, Any] = {}
    had_audio = False
    if request.content_type and "multipart/form-data" in request.content_type:
        transcript = (request.form.get("transcript") or "").strip()
        f = request.files.get("audio")
        if f and f.filename:
            had_audio = True
            suffix = Path(f.filename).suffix or ".wav"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                f.save(tmp.name)
                tmp_path = Path(tmp.name)
            try:
                audio_meta = _analyze_audio_file(tmp_path)
                recognized = _transcribe_audio_file(tmp_path)
                if recognized:
                    transcript = f"{transcript}\n{recognized}".strip()
            except Exception as e:
                audio_meta = {"error": str(e)}
            finally:
                try:
                    tmp_path.unlink(missing_ok=True)
                except Exception:
                    pass
    else:
        data = request.get_json() or {}
        transcript = (data.get("transcript") or "").strip()
    if not transcript and not audio_meta and not had_audio:
        return jsonify({"error": "transcript or audio required"}), 400
    if not transcript and (audio_meta or had_audio):
        transcript = "Аудио без распознанного текста."

    brain = get_brain()
    if brain and getattr(brain, "_llm", None):
        tl = transcript.lower()
        tempo_val = float(audio_meta.get("tempo") or 0) if isinstance(audio_meta, dict) else 0.0
        has_music_markers = any(k in tl for k in ("куплет", "припев", "бит", "мелод", "песня", "трек"))
        if has_music_markers or tempo_val > 70:
            mode_hint = "музыка/аудио"
        else:
            mode_hint = "речь/звуки"
        audio_meta_text = ""
        if audio_meta:
            audio_meta_text = (
                f"\nМетрики аудио: tempo={audio_meta.get('tempo')}, "
                f"energy={audio_meta.get('energy')}, mood_hint={audio_meta.get('mood_hint')}, "
                f"melody_hint={audio_meta.get('melody_hint')}, duration_sec={audio_meta.get('duration_sec')}"
            )
        lyrics_hint = transcript if transcript and transcript != "Аудио без распознанного текста." else "(текст песни не распознан)"
        prompt = (
            "Ты Даша. Ответь как живой собеседник и используй цепочку анализа:\n"
            "1) Что слышно в мелодии/звуке (темп, энергия, характер),\n"
            "2) Что слышно в словах/тексте,\n"
            "3) Общее ощущение и короткий вывод для пользователя.\n"
            "Важно: не выдумывай события/новости/сюжет, если этого нет в данных. "
            "Если данных мало, прямо скажи, что вывод ограничен.\n"
            f"Тип аудио: {mode_hint}.\n\n"
            f"Текст/транскрипт: {lyrics_hint}{audio_meta_text}"
        )
        try:
            r = brain._llm.generate([
                {"role": "system", "content": "Ты анализируешь услышанный текст как эмпатичный ассистент."},
                {"role": "user", "content": prompt},
            ])
            return jsonify({
                "status": "ok",
                "result": r.content,
                "audio_meta": audio_meta,
                "stages": {
                    "melody": {
                        "tempo": audio_meta.get("tempo"),
                        "energy": audio_meta.get("energy"),
                        "melody_hint": audio_meta.get("melody_hint"),
                    },
                    "lyrics": {"transcript": transcript},
                    "summary": {"mood_hint": audio_meta.get("mood_hint")},
                },
            })
        except Exception as e:
            return jsonify({"error": str(e)}), 500
    return jsonify({"status": "ok", "result": f"Я услышала: {transcript}", "audio_meta": audio_meta})


@app.route("/api/files")
def api_files_list():
    path = request.args.get("path", "")
    target = FILES_DIR / path
    if not target.exists() or not str(target.resolve()).startswith(str(FILES_DIR.resolve())):
        return jsonify({"error": "Invalid path"}), 400
    
    items = []
    for item in target.iterdir():
        stat = item.stat()
        items.append({
            "name": item.name,
            "path": str(item.relative_to(FILES_DIR)),
            "is_dir": item.is_dir(),
            "size": stat.st_size if item.is_file() else 0,
            "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        })
    items.sort(key=lambda x: (not x["is_dir"], x["name"].lower()))
    return jsonify({"path": path, "items": items})


@app.route("/api/files/read")
def api_files_read():
    path = request.args.get("path", "")
    target = (FILES_DIR / path).resolve()
    if not str(target).startswith(str(FILES_DIR.resolve())):
        return jsonify({"error": "Invalid path"}), 400
    if not target.exists() or not target.is_file():
        return jsonify({"error": "Not found"}), 404
    try:
        content = _read_file_content(target)
        return jsonify({"content": content, "ext": target.suffix.lower()})
    except Exception as e:
        return jsonify({"error": f"Cannot read: {e}"}), 400


@app.route("/api/files/write", methods=["POST"])
def api_files_write():
    data = request.get_json() or {}
    path, content = data.get("path", ""), data.get("content", "")
    if not path:
        return jsonify({"error": "Path required"}), 400
    target = (FILES_DIR / path).resolve()
    if not str(target).startswith(str(FILES_DIR.resolve())):
        return jsonify({"error": "Invalid path"}), 400
    if _is_diary_protected_target(target):
        return jsonify({"error": "Diary is read-only for user edits"}), 403
    target.parent.mkdir(parents=True, exist_ok=True)
    try:
        _write_file_content(target, content)
        return jsonify({"status": "ok"})
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@app.route("/api/files/apply-assist", methods=["POST"])
def api_files_apply_assist():
    data = request.get_json() or {}
    path = data.get("path", "")
    content = data.get("content", "")
    if not path:
        return jsonify({"error": "Path required"}), 400
    target = (FILES_DIR / path).resolve()
    if not str(target).startswith(str(FILES_DIR.resolve())):
        return jsonify({"error": "Invalid path"}), 400
    if _is_diary_protected_target(target):
        return jsonify({"error": "Diary is read-only for user edits"}), 403
    target.parent.mkdir(parents=True, exist_ok=True)
    try:
        _write_file_content(target, content)
        return jsonify({"status": "ok"})
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@app.route("/api/files/mkdir", methods=["POST"])
def api_files_mkdir():
    path = (request.get_json() or {}).get("path", "")
    if path:
        target = (FILES_DIR / path).resolve()
        if not str(target).startswith(str(FILES_DIR.resolve())):
            return jsonify({"error": "Invalid path"}), 400
        if _is_diary_protected_target(target):
            return jsonify({"error": "Diary is read-only for user edits"}), 403
        target.mkdir(parents=True, exist_ok=True)
    return jsonify({"status": "ok"})


@app.route("/api/files/delete", methods=["POST"])
def api_files_delete():
    path = (request.get_json() or {}).get("path", "")
    target = (FILES_DIR / path).resolve()
    if not str(target).startswith(str(FILES_DIR.resolve())):
        return jsonify({"error": "Invalid path"}), 400
    if _is_diary_protected_target(target):
        return jsonify({"error": "Diary is read-only for user edits"}), 403
    if target.is_dir():
        shutil.rmtree(target)
    elif target.is_file():
        target.unlink()
    return jsonify({"status": "ok"})


@app.route("/api/files/upload", methods=["POST"])
def api_files_upload():
    if 'file' not in request.files:
        return jsonify({"error": "No file"}), 400
    file = request.files['file']
    path = request.form.get("path", "")
    target_dir = (FILES_DIR / path).resolve() if path else FILES_DIR.resolve()
    if not str(target_dir).startswith(str(FILES_DIR.resolve())):
        return jsonify({"error": "Invalid path"}), 400
    if _is_diary_protected_target(target_dir):
        return jsonify({"error": "Diary is read-only for user edits"}), 403
    target_dir.mkdir(parents=True, exist_ok=True)
    file.save(target_dir / file.filename)
    return jsonify({"status": "ok"})


@app.route("/api/files/download/<path:filepath>")
def api_files_download(filepath):
    return send_from_directory(FILES_DIR, filepath, as_attachment=True)


# ═══════════════════════════════════════════════════════════════════
#  Logs & Notifications
# ═══════════════════════════════════════════════════════════════════

@app.route("/api/logs")
def api_logs():
    return jsonify(web_log_handler.get_logs(request.args.get("limit", 100, type=int)))


@app.route("/api/logs/stream")
def api_logs_stream():
    def generate():
        q = web_log_handler.subscribe()
        try:
            while True:
                try:
                    yield f"data: {json.dumps(q.get(timeout=30))}\n\n"
                except queue.Empty:
                    yield ": keepalive\n\n"
        except (GeneratorExit, OSError, ConnectionError):
            pass
        finally:
            web_log_handler.unsubscribe(q)
    return Response(
        generate(),
        mimetype='text/event-stream',
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


@app.route("/api/notifications")
def api_notifications():
    return jsonify(notifications.get_all())


@app.route("/api/notifications/add", methods=["POST"])
def api_notifications_add():
    data = request.get_json() or {}
    return jsonify(notifications.add(
        data.get("title", ""), data.get("message", ""),
        data.get("type", "info"), data.get("icon", "💬"),
        data.get("duration", 5000), data.get("action"), data.get("action_data"),
        bool(data.get("system", False))
    ))


@app.route("/api/notifications/stream")
def api_notifications_stream():
    def generate():
        q = notifications.subscribe()
        try:
            while True:
                try:
                    yield f"data: {json.dumps(q.get(timeout=30))}\n\n"
                except queue.Empty:
                    yield ": keepalive\n\n"
        except (GeneratorExit, OSError, ConnectionError):
            pass
        finally:
            notifications.unsubscribe(q)
    return Response(
        generate(),
        mimetype='text/event-stream',
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


@app.route("/api/system/info")
def api_system_info():
    try:
        import psutil
        return jsonify({
            "cpu_percent": psutil.cpu_percent(),
            "memory_percent": psutil.virtual_memory().percent,
            "disk_percent": psutil.disk_usage('/').percent,
        })
    except:
        return jsonify({})


@app.route("/api/daria/metrics")
def api_daria_metrics():
    """Process-level metrics for Daria monitor window."""
    metrics = {
        "uptime_sec": max(0.0, time.time() - PROCESS_START_TS),
        "threads": threading.active_count(),
        "notifications_total": len(notifications.get_all(2000)),
        "tasks_today": len(task_manager.list_all().get("dasha_tasks", [])),
    }
    try:
        import psutil  # type: ignore
        p = psutil.Process(os.getpid())
        with p.oneshot():
            mem = p.memory_info()
            metrics.update({
                "cpu_percent": p.cpu_percent(interval=0.0),
                "rss_mb": round(mem.rss / (1024 * 1024), 2),
                "vms_mb": round(mem.vms / (1024 * 1024), 2),
                "open_files": len(p.open_files() or []),
            })
    except Exception:
        metrics.update({
            "cpu_percent": 0.0,
            "rss_mb": 0.0,
            "vms_mb": 0.0,
            "open_files": 0,
        })
    return jsonify(metrics)


@app.route("/api/debug/runtime")
def api_debug_runtime():
    """Detailed runtime diagnostics for debug window."""
    plugins_data = []
    try:
        pm = get_plugins()
        if pm:
            for s in pm.get_installed_plugins():
                plugins_data.append({
                    "id": s.manifest.id,
                    "name": s.manifest.name,
                    "enabled": bool(s.enabled),
                    "loaded": bool(s.loaded),
                })
    except Exception:
        plugins_data = []

    threads = []
    for t in threading.enumerate():
        threads.append({
            "name": t.name,
            "daemon": bool(t.daemon),
            "alive": bool(t.is_alive()),
        })

    recent = web_log_handler.get_logs(500)
    errors = [x for x in recent if x.get("level") in ("ERROR", "CRITICAL")]
    warnings = [x for x in recent if x.get("level") == "WARNING"]
    req_chain = [x for x in recent if "REQ[" in str(x.get("message", "")) or "RES[" in str(x.get("message", ""))]
    model_map = {
        "chat_llm": "Ollama (core.llm)",
        "vision": str(load_settings().get("senses_vision_provider", "auto")),
        "audio_asr": str(load_settings().get("senses_audio_provider", "auto")),
        "image_gen": str(load_settings().get("image_gen_model", "Tongyi-MAI/Z-Image-Turbo")),
        "music_analysis": "librosa",
    }
    jobs = []
    try:
        jobs = sorted(list(image_jobs.jobs.values()), key=lambda x: x.get("updated_at", ""), reverse=True)[:20]
    except Exception:
        jobs = []
    return jsonify({
        "status": "ok",
        "process": {
            "pid": os.getpid(),
            "uptime_sec": max(0.0, time.time() - PROCESS_START_TS),
            "python": os.sys.version.split()[0],
        },
        "threads": threads,
        "plugins": plugins_data,
        "errors_last": errors[-25:],
        "warnings_last": warnings[-25:],
        "requests_last": req_chain[-60:],
        "image_jobs_last": jobs,
        "models": model_map,
    })


# ═══════════════════════════════════════════════════════════════════
#  Plugins
# ═══════════════════════════════════════════════════════════════════

@app.route("/api/plugins")
def api_plugins_list():
    plugins = get_plugins()
    if not plugins:
        return jsonify([])
    return jsonify([{**s.manifest.to_dict(), "installed": True, "enabled": s.enabled, "loaded": s.loaded} 
                    for s in plugins.get_installed_plugins()])


@app.route("/api/plugins/desktop")
def api_plugins_desktop():
    plugins = get_plugins()
    return jsonify(plugins.get_desktop_plugins() if plugins else [])


@app.route("/api/plugins/catalog")
def api_plugins_catalog():
    plugins = get_plugins()
    return jsonify(plugins.fetch_catalog() if plugins else [])


@app.route("/api/plugins/<plugin_id>")
def api_plugin_info(plugin_id):
    plugins = get_plugins()
    if not plugins:
        return jsonify({"error": "unavailable"}), 500
    info = plugins.get_plugin_info(plugin_id)
    if info:
        for upd in plugins.check_plugin_updates():
            if upd.get("id") == plugin_id:
                info["update_available"] = True
                info["latest_version"] = upd.get("latest_version")
                break
        else:
            info["update_available"] = False
            info["latest_version"] = info.get("version")
        return jsonify(info)
    for item in plugins.fetch_catalog():
        if item.get("id") == plugin_id:
            return jsonify(item)
    return jsonify({"error": "Not found"}), 404


@app.route("/api/plugins/<plugin_id>/install", methods=["POST"])
def api_plugin_install(plugin_id):
    plugins = get_plugins()
    if plugins and plugins.install_plugin(plugin_id):
        notifications.add("Плагин", f"{plugin_id} установлен", "success", "🧩")
        return jsonify({"status": "ok"})
    return jsonify({"error": "Failed"}), 500


@app.route("/api/plugins/updates")
def api_plugins_updates():
    plugins = get_plugins()
    if not plugins:
        return jsonify([])
    try:
        return jsonify(plugins.check_plugin_updates())
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/plugins/<plugin_id>/update", methods=["POST"])
def api_plugin_update(plugin_id):
    plugins = get_plugins()
    if plugins and plugins.update_plugin(plugin_id):
        notifications.add("Плагин", f"{plugin_id} обновлён", "success", "🧩")
        return jsonify({"status": "ok"})
    return jsonify({"error": "Failed"}), 500


@app.route("/api/plugins/update-all", methods=["POST"])
def api_plugins_update_all():
    plugins = get_plugins()
    if not plugins:
        return jsonify({"error": "unavailable"}), 500
    try:
        return jsonify({"status": "ok", **plugins.update_all_plugins()})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/plugins/<plugin_id>/uninstall", methods=["POST"])
def api_plugin_uninstall(plugin_id):
    plugins = get_plugins()
    if plugins and plugins.uninstall_plugin(plugin_id):
        return jsonify({"status": "ok"})
    return jsonify({"error": "Failed"}), 500


@app.route("/api/plugins/<plugin_id>/window")
def api_plugin_window(plugin_id):
    plugins = get_plugins()
    return jsonify(plugins.get_plugin_window_data(plugin_id) if plugins else {})


@app.route("/api/plugins/<plugin_id>/action", methods=["POST"])
def api_plugin_action(plugin_id):
    plugins = get_plugins()
    if not plugins:
        return jsonify({"error": "unavailable"}), 500
    data = request.get_json() or {}
    return jsonify(plugins.call_plugin_action(plugin_id, data.get("action", ""), data.get("data", {})))


@app.route("/plugins/<plugin_id>/static/<path:filename>")
def plugin_static(plugin_id, filename):
    from core.config import get_config
    return send_from_directory(get_config().data_dir / "plugins" / plugin_id / "static", filename)


@app.route("/plugins/<plugin_id>/template/<filename>")
def plugin_template(plugin_id, filename):
    from core.config import get_config
    path = get_config().data_dir / "plugins" / plugin_id / "templates" / filename
    return path.read_text() if path.exists() else abort(404)


# ═══════════════════════════════════════════════════════════════════
#  Updater
# ═══════════════════════════════════════════════════════════════════

@app.route("/api/update/check")
def api_update_check():
    source = request.args.get("source", "github")
    repo = request.args.get("repo", "dariumi/Daria")
    ref = request.args.get("ref", "main")
    current = _read_version(PROJECT_ROOT / "VERSION")
    latest = current
    update_available = False

    if source == "github":
        try:
            try:
                import requests
                api_url = f"https://api.github.com/repos/{repo.strip('/')}/releases/latest"
                r = requests.get(api_url, timeout=15)
                if r.status_code == 200:
                    latest = str((r.json().get("tag_name") or "").lstrip("v") or current)
                else:
                    tags_url = f"https://api.github.com/repos/{repo.strip('/')}/tags"
                    r2 = requests.get(tags_url, timeout=15)
                    if r2.status_code == 200 and r2.json():
                        latest = str(r2.json()[0].get("name", "")).lstrip("v") or current
            except Exception:
                latest = current
            update_available = _version_key(latest) > _version_key(current)
        except Exception as e:
            return jsonify({"error": str(e), "current": current, "latest": current, "update_available": False}), 200

    _update_state["last_check"] = datetime.now().isoformat()
    return jsonify({
        "current": current,
        "latest": latest,
        "update_available": update_available,
        "source": source,
        "repo": repo,
        "ref": ref,
    })


@app.route("/api/update/auto", methods=["GET", "POST"])
def api_update_auto():
    settings = load_settings()
    if request.method == "GET":
        return jsonify({"auto_update": settings.get("auto_update", False)})
    data = request.get_json() or {}
    enabled = bool(data.get("auto_update", False))
    save_settings({"auto_update": enabled})
    return jsonify({"status": "ok", "auto_update": enabled})


@app.route("/api/update/from-github", methods=["POST"])
def api_update_from_github():
    if _update_state.get("running"):
        return jsonify({"error": "update already running"}), 409
    data = request.get_json() or {}
    repo = data.get("repo", "dariumi/Daria")
    ref = data.get("ref", "main")
    _update_state["running"] = True
    _update_state["last_error"] = None
    _update_state["last_action"] = f"github:{repo}@{ref}"
    try:
        archive = _download_github_archive(repo, ref)
        with tempfile.TemporaryDirectory(prefix="daria-update-") as td:
            archive_path = Path(td) / "update.zip"
            archive_path.write_bytes(archive)
            extracted = Path(td) / "extracted"
            _extract_archive_to_dir(archive_path, extracted)
            src_root = _find_project_root(extracted)
            _sync_project_tree(src_root, PROJECT_ROOT)
        new_version = _read_version(PROJECT_ROOT / "VERSION")
        notifications.add("Обновление", f"Система обновлена до v{new_version}", "success", "⬆️", 10000)
        return jsonify({"status": "ok", "version": new_version, "restart_required": True})
    except Exception as e:
        _update_state["last_error"] = str(e)
        return jsonify({"error": str(e)}), 500
    finally:
        _update_state["running"] = False


@app.route("/api/update/from-archive", methods=["POST"])
def api_update_from_archive():
    if _update_state.get("running"):
        return jsonify({"error": "update already running"}), 409
    data = request.get_json() or {}
    archive_path = Path((data.get("archive_path") or "").strip()).expanduser()
    if not archive_path.exists():
        return jsonify({"error": "archive not found"}), 400

    _update_state["running"] = True
    _update_state["last_error"] = None
    _update_state["last_action"] = f"archive:{archive_path}"
    try:
        with tempfile.TemporaryDirectory(prefix="daria-update-") as td:
            extracted = Path(td) / "extracted"
            _extract_archive_to_dir(archive_path, extracted)
            src_root = _find_project_root(extracted)
            _sync_project_tree(src_root, PROJECT_ROOT)
        new_version = _read_version(PROJECT_ROOT / "VERSION")
        notifications.add("Обновление", f"Система обновлена до v{new_version}", "success", "⬆️", 10000)
        return jsonify({"status": "ok", "version": new_version, "restart_required": True})
    except Exception as e:
        _update_state["last_error"] = str(e)
        return jsonify({"error": str(e)}), 500
    finally:
        _update_state["running"] = False


@app.route("/api/update/state")
def api_update_state():
    return jsonify({
        **_update_state,
        "version": _read_version(PROJECT_ROOT / "VERSION"),
    })


# ═══════════════════════════════════════════════════════════════════
#  Wiki
# ═══════════════════════════════════════════════════════════════════

@app.route("/api/wiki/pages")
def api_wiki_pages():
    wiki_dir = PROJECT_ROOT / "docs" / "wiki"
    wiki_dir.mkdir(parents=True, exist_ok=True)
    pages = sorted([p.name for p in wiki_dir.glob("*.md")])
    return jsonify({"pages": pages})


@app.route("/api/wiki/page")
def api_wiki_page():
    name = (request.args.get("name") or "Home.md").strip()
    if "/" in name or "\\" in name:
        return jsonify({"error": "Invalid page"}), 400
    wiki_dir = PROJECT_ROOT / "docs" / "wiki"
    path = (wiki_dir / name).resolve()
    if not path.exists() or not str(path).startswith(str(wiki_dir.resolve())):
        return jsonify({"error": "Not found"}), 404
    text = path.read_text(encoding="utf-8")
    return jsonify({"name": name, "content": text})


@app.route("/api/browser/proxy")
def api_browser_proxy():
    raw = (request.args.get("url") or "").strip()
    if not raw:
        return "<h3>URL не указан</h3>", 400
    if not re.match(r"^https?://", raw, flags=re.IGNORECASE):
        raw = "https://" + raw
    try:
        parsed = urllib.parse.urlparse(raw)
        if parsed.scheme not in ("http", "https"):
            return "<h3>Поддерживаются только http/https</h3>", 400
        # Normalize URL to avoid ascii codec issues for Cyrillic query/path.
        safe_path = urllib.parse.quote(parsed.path or "/", safe="/%:@+~!$&'()*;,=-._")
        safe_query = urllib.parse.quote_plus(parsed.query, safe="=&:%@+~!$'()*;,.-_")
        safe_url = urllib.parse.urlunparse((
            parsed.scheme,
            parsed.netloc.encode("idna").decode("ascii"),
            safe_path,
            parsed.params,
            safe_query,
            parsed.fragment,
        ))
        req = urllib.request.Request(
            safe_url,
            headers={
                "User-Agent": "DARIA-Browser/0.9.1",
                "Accept-Language": "ru,en;q=0.8",
            },
        )
        with urllib.request.urlopen(req, timeout=12) as resp:
            ctype = (resp.headers.get("Content-Type") or "").lower()
            data = resp.read()
        if "text/html" in ctype or not ctype:
            html_text = data.decode("utf-8", errors="ignore")
            if "<head" in html_text.lower():
                html_text = re.sub(
                    r"(?i)<head([^>]*)>",
                    r"<head\1><base href=\"" + safe_url + "\"><style>body{max-width:1200px;margin:0 auto;padding:12px;font-family:Arial,sans-serif}</style>",
                    html_text,
                    count=1,
                )
            return Response(html_text, mimetype="text/html")
        if ctype.startswith("image/"):
            return Response(data, mimetype=ctype.split(";")[0])
        text = data.decode("utf-8", errors="ignore")
        return Response(f"<pre>{html.escape(text)}</pre>", mimetype="text/html")
    except Exception as e:
        return Response(
            f"<h3>Не удалось открыть страницу</h3><p>{html.escape(str(e))}</p>"
            f"<p><a href=\"{html.escape(raw)}\" target=\"_blank\" rel=\"noopener\">Открыть напрямую</a></p>",
            mimetype="text/html",
            status=502,
        )


@app.route("/api/browser/start")
def api_browser_start():
    art = (
        "data:image/svg+xml;utf8,"
        + urllib.parse.quote(
            "<svg xmlns='http://www.w3.org/2000/svg' width='1200' height='520'>"
            "<defs><linearGradient id='g' x1='0' y1='0' x2='1' y2='1'>"
            "<stop offset='0%' stop-color='#f472b6'/><stop offset='100%' stop-color='#60a5fa'/></linearGradient></defs>"
            "<rect width='100%' height='100%' fill='#0b1220'/>"
            "<circle cx='240' cy='120' r='210' fill='url(#g)' fill-opacity='0.25'/>"
            "<circle cx='980' cy='420' r='260' fill='url(#g)' fill-opacity='0.22'/>"
            "<text x='80' y='220' fill='#fff' font-size='54' font-family='Arial'>DARIA Browser</text>"
            "<text x='82' y='270' fill='#cbd5e1' font-size='24' font-family='Arial'>Поиск, чтение и наблюдение вместе с Дашей</text>"
            "</svg>"
        )
    )
    html_text = (
        "<!doctype html><html><head><meta charset='utf-8'><title>DARIA Browser</title>"
        "<style>body{margin:0;background:#0b1220;color:#fff;font-family:Arial,sans-serif}"
        ".wrap{padding:24px;max-width:1060px;margin:0 auto}.hero{width:100%;border-radius:16px;border:1px solid #334155}"
        ".hint{margin-top:16px;color:#cbd5e1;font-size:15px}</style></head><body>"
        "<div class='wrap'><img class='hero' alt='DARIA' src='" + art + "'/>"
        "<p class='hint'>Введи запрос или URL в верхней строке. Поддерживается поиск на русскоязычных сайтах.</p>"
        "</div></body></html>"
    )
    return Response(html_text, mimetype="text/html")


@app.route("/wiki")
def wiki_redirect():
    return render_template("index.html", version=VERSION)


# ═══════════════════════════════════════════════════════════════════
#  Server
# ═══════════════════════════════════════════════════════════════════

def create_app():
    return app


def run_server(host: str = "127.0.0.1", port: int = 8000, 
               debug: bool = False, ssl_context = None):
    logger.info("Initializing DARIA...")
    ensure_sample_books()
    get_brain()
    get_memory()
    get_plugins()
    
    # Start attention thread
    settings = load_settings()
    # Do not preload or warm up image/vision models on server startup.
    # Models are loaded lazily only when a related request arrives.
    attention_thread.enabled = settings.get("attention_enabled", True)
    if not attention_thread.is_alive():
        attention_thread.start()
    if not activity_thread.is_alive():
        activity_thread.start()

    # Treat offline period as sleep/rest for more human-like rhythm.
    try:
        lifecycle_file = DATA_DIR / "lifecycle.json"
        now = datetime.now()
        offline_hours = 0.0
        if lifecycle_file.exists():
            prev = json.loads(lifecycle_file.read_text(encoding="utf-8"))
            last_shutdown = prev.get("last_shutdown")
            if last_shutdown:
                dt = datetime.fromisoformat(last_shutdown)
                offline_hours = max(0.0, (now - dt).total_seconds() / 3600.0)
        brain = get_brain()
        if brain and offline_hours > 0 and hasattr(brain, "mood"):
            if 0.3 <= offline_hours <= 2.5:
                brain.mood.energy = min(1.0, brain.mood.energy + 0.2)
                notifications.add("🌸 Даша", "Кажется, я немного вздремнула и стала бодрее.", "info", "😴", 7000)
            elif offline_hours > 8:
                brain.mood.energy = min(1.0, brain.mood.energy + 0.35)
                notifications.add("🌸 Даша", "Я проснулась после долгого сна, доброе возвращение.", "info", "🌅", 7000)
            else:
                brain.mood.energy = max(0.35, brain.mood.energy)
        lifecycle_file.write_text(json.dumps({"last_start": now.isoformat()}, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as e:
        logger.debug(f"Lifecycle init error: {e}")
    
    notifications.add("DARIA", f"Система запущена v{VERSION}", "success", "🌸", 8000)
    logger.info("Ready!")
    
    # Keep single-process runtime even in debug mode:
    # Werkzeug reloader spawns a parent+child process, which duplicates
    # background workers/plugins (e.g. Telegram polling -> 409 Conflict).
    app.run(
        host=host,
        port=port,
        debug=debug,
        threaded=True,
        ssl_context=ssl_context,
        use_reloader=False,
    )


application = app
